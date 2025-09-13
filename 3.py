# EFL Café Question Generator — Full Single-File App (HiDPI-safe exact preview on demand)
# - PyQt5 UI with optional exact DOCX-styled preview (via button)
# - Updater lives in About… dialog, shows version info
# - Gemini 2.5 Flash integration (batch + single regen, feedback.txt logging)
# - Editable topics, level, and questions (double-click to edit rows)
# - Export to DOCX using your template
# - DPI-aware; starts full screen
# - Robust, drop-in-friendly exact preview subsystem (DOCX→PDF→PNG)

import os, sys, re, json, hashlib, tempfile, shutil, subprocess, traceback
from urllib.request import Request, urlopen
from urllib.error import URLError, HTTPError

# Debug mode detection
DEBUG_MODE = False
if getattr(sys, 'frozen', False):
    # Running as compiled executable
    exe_dir = os.path.dirname(sys.executable)
    debug_file = os.path.join(exe_dir, "DEBUG.txt")
    DEBUG_MODE = os.path.exists(debug_file)
else:
    # Running as Python script
    DEBUG_MODE = True

# Redirect stdout and stderr to prevent console flashes in compiled exe
if getattr(sys, 'frozen', False) and not DEBUG_MODE:
    # Redirect stdout and stderr to null to prevent console window
    import os
    devnull = open(os.devnull, 'w')
    sys.stdout = devnull
    sys.stderr = devnull

def debug_print(message):
    """Print debug messages only in debug mode to avoid console flashes."""
    if DEBUG_MODE:
        print(message)

# OCR imports
try:
    from PIL import Image, ImageOps
    import pytesseract
    import fitz  # PyMuPDF
    OCR_AVAILABLE = True
except ImportError as e:
    OCR_AVAILABLE = False
    debug_print(f"OCR dependencies not available: {e}")

# Google AI imports
try:
    import google.generativeai as genai
    GOOGLE_AI_AVAILABLE = True
    debug_print(f"Google AI library loaded successfully")
except ImportError as e:
    GOOGLE_AI_AVAILABLE = False
    debug_print(f"Google AI library not available: {e}")

# -------------------------------------------------------------------
# Enable Windows per-monitor DPI awareness BEFORE Qt starts
# -------------------------------------------------------------------
def _enable_win_per_monitor_dpi_awareness():
    if os.name != "nt":
        return
    import ctypes
    try:
        # Windows 10+ : Per-Monitor-V2 (best)
        ctypes.windll.user32.SetProcessDpiAwarenessContext(ctypes.c_void_p(-4))  # PER_MONITOR_AWARE_V2
        return
    except Exception:
        pass
    try:
        # Windows 8.1 : Per-Monitor (fallback)
        PROCESS_PER_MONITOR_DPI_AWARE = 2
        ctypes.windll.shcore.SetProcessDpiAwareness(PROCESS_PER_MONITOR_DPI_AWARE)
        return
    except Exception:
        pass
    try:
        # Vista/7 : system-DPI aware (last resort)
        ctypes.windll.user32.SetProcessDPIAware()
    except Exception:
        pass

# Environment flags must be set before QApplication is created
os.environ.setdefault("QT_ENABLE_HIGHDPI_SCALING", "1")
os.environ.setdefault("QT_SCALE_FACTOR_ROUNDING_POLICY", "PassThrough")

# --- PyQt imports (after setting env) ---
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer, QSize, QRect, QEvent, QCoreApplication
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import (
    QApplication, QWidget, QMessageBox, QDialog, QVBoxLayout, QLabel,
    QProgressBar, QPushButton, QHBoxLayout, QVBoxLayout as QVLayout, QSplitter,
    QScrollArea, QFrame, QProgressDialog, QFileDialog, QTextEdit, QLineEdit,
    QComboBox, QSizePolicy, QInputDialog, QRadioButton, QButtonGroup, QCheckBox,
    QGridLayout
)
from PyQt5.QtGui import QFont, QPixmap, QPainter, QColor

# =========================
# App Info & Config
# =========================
APP_NAME    = "EFL Cafe Wizard"
APP_VERSION = "0.23 BETA"  # updated for beta release
UPDATE_IN_PROGRESS = False  # Global flag to prevent multiple update dialogs

# =========================
# Standard Button Styles
# =========================
def get_standard_button_style(button_type="primary"):
    """Get standardized button styles for the application."""
    styles = {
        "primary": """
            QPushButton {
                background-color: #007acc;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 9pt;
            }
            QPushButton:hover {
                background-color: #005a9e;
            }
            QPushButton:pressed {
                background-color: #004578;
            }
            QPushButton:disabled {
                background-color: #6c757d;
                color: #adb5bd;
            }
        """,
        "secondary": """
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 9pt;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
            QPushButton:pressed {
                background-color: #495057;
            }
            QPushButton:disabled {
                background-color: #adb5bd;
                color: #6c757d;
            }
        """,
        "success": """
            QPushButton {
                background-color: #28a745;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 9pt;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1e7e34;
            }
            QPushButton:disabled {
                background-color: #6c757d;
                color: #adb5bd;
            }
        """,
        "warning": """
            QPushButton {
                background-color: #ffc107;
                color: #212529;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 9pt;
            }
            QPushButton:hover {
                background-color: #e0a800;
            }
            QPushButton:pressed {
                background-color: #d39e00;
            }
            QPushButton:disabled {
                background-color: #6c757d;
                color: #adb5bd;
            }
        """,
        "danger": """
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 9pt;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
            QPushButton:pressed {
                background-color: #bd2130;
            }
            QPushButton:disabled {
                background-color: #6c757d;
                color: #adb5bd;
            }
        """,
        "outline": """
            QPushButton {
                background-color: transparent;
                color: #007acc;
                border: 1px solid #007acc;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: bold;
                font-size: 9pt;
            }
            QPushButton:hover {
                background-color: #007acc;
                color: white;
            }
            QPushButton:pressed {
                background-color: #005a9e;
                color: white;
            }
            QPushButton:disabled {
                background-color: transparent;
                color: #6c757d;
                border-color: #6c757d;
            }
        """,
        "minimal": """
            QPushButton {
                background-color: #e9ecef;
                color: #495057;
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 6px 12px;
                font-weight: normal;
                font-size: 10pt;
            }
            QPushButton:hover {
                background-color: #dee2e6;
            }
            QPushButton:pressed {
                background-color: #ced4da;
            }
            QPushButton:disabled {
                background-color: #f8f9fa;
                color: #6c757d;
                border-color: #e9ecef;
            }
        """,
        "small": """
            QPushButton {
                background-color: transparent;
                color: #007acc;
                border: 1px solid #007acc;
                border-radius: 3px;
                padding: 4px 8px;
                font-weight: bold;
                font-size: 9pt;
                min-width: 30px;
            }
            QPushButton:hover {
                background-color: #007acc;
                color: white;
            }
            QPushButton:pressed {
                background-color: #005a9e;
                color: white;
            }
            QPushButton:disabled {
                background-color: transparent;
                color: #6c757d;
                border-color: #6c757d;
            }
        """,
        "dialog": """
            QPushButton {
                background-color: #007acc;
                color: white;
                border: none;
                border-radius: 3px;
                padding: 6px 12px;
                font-weight: bold;
                font-size: 10pt;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #005a9e;
            }
            QPushButton:pressed {
                background-color: #004578;
            }
            QPushButton:disabled {
                background-color: #6c757d;
                color: #adb5bd;
            }
        """,
        "dialog_secondary": """
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: none;
                border-radius: 3px;
                padding: 6px 12px;
                font-weight: bold;
                font-size: 10pt;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
            QPushButton:pressed {
                background-color: #495057;
            }
            QPushButton:disabled {
                background-color: #adb5bd;
                color: #6c757d;
            }
        """,
        "dialog_outline": """
            QPushButton {
                background-color: transparent;
                color: #007acc;
                border: 1px solid #007acc;
                border-radius: 3px;
                padding: 6px 12px;
                font-weight: bold;
                font-size: 10pt;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #007acc;
                color: white;
            }
            QPushButton:pressed {
                background-color: #005a9e;
                color: white;
            }
            QPushButton:disabled {
                background-color: transparent;
                color: #6c757d;
                border-color: #6c757d;
            }
        """
    }
    return styles.get(button_type, styles["primary"])

def get_standard_combo_style():
    """Get standardized combo box styles to match button themes."""
    return """
        QComboBox {
            background-color: white;
            color: #495057;
            border: 1px solid #ced4da;
            border-radius: 4px;
            padding: 6px 12px;
            font-size: 11pt;
            min-width: 100px;
        }
        QComboBox:hover {
            border-color: #007acc;
        }
        QComboBox:focus {
            border-color: #007acc;
            outline: none;
        }
        QComboBox::drop-down {
            border: none;
            width: 20px;
        }
        QComboBox::down-arrow {
            image: none;
            border-left: 5px solid transparent;
            border-right: 5px solid transparent;
            border-top: 5px solid #6c757d;
            margin-right: 5px;
        }
        QComboBox QAbstractItemView {
            background-color: white;
            border: 1px solid #ced4da;
            border-radius: 4px;
            selection-background-color: #007acc;
            selection-color: white;
            padding: 4px;
        }
        QComboBox QAbstractItemView::item {
            padding: 6px 12px;
            min-height: 20px;
        }
        QComboBox QAbstractItemView::item:hover {
            background-color: #f8f9fa;
        }
    """

def get_standard_checkbox_style():
    """Get standardized checkbox styles to match button themes."""
    return """
        QCheckBox {
            color: #495057;
            font-size: 11pt;
            spacing: 8px;
        }
        QCheckBox::indicator {
            width: 16px;
            height: 16px;
            border: 1px solid #ced4da;
            border-radius: 3px;
            background-color: white;
        }
        QCheckBox::indicator:hover {
            border-color: #007acc;
        }
        QCheckBox::indicator:checked {
            background-color: #007acc;
            border-color: #007acc;
            image: url(data:image/svg+xml;base64,PHN2ZyB3aWR0aD0iMTIiIGhlaWdodD0iOSIgdmlld0JveD0iMCAwIDEyIDkiIGZpbGw9Im5vbmUiIHhtbG5zPSJodHRwOi8vd3d3LnczLm9yZy8yMDAwL3N2ZyI+CjxwYXRoIGQ9Ik0xIDQuNUw0LjUgOEwxMSAxIiBzdHJva2U9IndoaXRlIiBzdHJva2Utd2lkdGg9IjIiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIgc3Ryb2tlLWxpbmVqb2luPSJyb3VuZCIvPgo8L3N2Zz4K);
        }
        QCheckBox::indicator:checked:hover {
            background-color: #005a9e;
            border-color: #005a9e;
        }
        QCheckBox:disabled {
            color: #6c757d;
        }
        QCheckBox::indicator:disabled {
            background-color: #f8f9fa;
            border-color: #e9ecef;
        }
    """

def get_standard_radio_style():
    """Get standardized radio button styles to match button themes."""
    return """
        QRadioButton {
            color: #495057;
            font-size: 11pt;
            spacing: 8px;
        }
        QRadioButton::indicator {
            width: 16px;
            height: 16px;
            border: 1px solid #ced4da;
            border-radius: 8px;
            background-color: white;
        }
        QRadioButton::indicator:hover {
            border-color: #007acc;
        }
        QRadioButton::indicator:checked {
            background-color: #007acc;
            border-color: #007acc;
        }
        QRadioButton::indicator:checked:hover {
            background-color: #005a9e;
            border-color: #005a9e;
        }
        QRadioButton:disabled {
            color: #6c757d;
        }
        QRadioButton::indicator:disabled {
            background-color: #f8f9fa;
            border-color: #e9ecef;
        }
    """

def get_standard_textedit_style():
    """Get standardized text edit styles to match button themes."""
    return """
        QTextEdit {
            background-color: white;
            color: #495057;
            border: 1px solid #ced4da;
            border-radius: 4px;
            padding: 8px;
            font-size: 11pt;
            selection-background-color: #007acc;
            selection-color: white;
        }
        QTextEdit:hover {
            border-color: #007acc;
        }
        QTextEdit:focus {
            border-color: #007acc;
            outline: none;
        }
        QTextEdit:disabled {
            background-color: #f8f9fa;
            color: #6c757d;
            border-color: #e9ecef;
        }
    """

def get_standard_lineedit_style():
    """Get standardized line edit styles to match button themes."""
    return """
        QLineEdit {
            background-color: white;
            color: #495057;
            border: 1px solid #ced4da;
            border-radius: 4px;
            padding: 6px 12px;
            font-size: 11pt;
            selection-background-color: #007acc;
            selection-color: white;
        }
        QLineEdit:hover {
            border-color: #007acc;
        }
        QLineEdit:focus {
            border-color: #007acc;
            outline: none;
        }
        QLineEdit:disabled {
            background-color: #f8f9fa;
            color: #6c757d;
            border-color: #e9ecef;
        }
    """

# Manifest-based updater (runs from About… dialog)
UPDATE_MANIFEST_URL = "https://raw.githubusercontent.com/aerladis/eflai/main/manifest.json"
NETWORK_TIMEOUT_SEC = 20

# Auto-update settings
AUTO_UPDATE_ENABLED = True
AUTO_UPDATE_INTERVAL_MS = 30 * 60 * 1000  # 30 minutes in milliseconds
AUTO_UPDATE_CHECK_ON_STARTUP = True

# Paths & models
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip() or "AIzaSyByCj3jsjlPaSZGsCqeDqR2filTFZ0abUQ"
GEMINI_MODEL   = "gemini-2.5-flash"

def _resource_path(rel: str) -> str:
    """Return absolute path to resource, works for dev and PyInstaller."""
    base = getattr(sys, "_MEIPASS", None)
    if base:
        return os.path.join(base, rel)
    here = os.path.dirname(__file__) or "."
    return os.path.join(here, rel)

def _get_prompts_path():
    """Get prompts.ini path, enabling debug mode only if DEBUG.txt exists."""
    debug_mode = False
    prompts_path = None
    
    # Check if we're running from PyInstaller bundle
    if getattr(sys, "frozen", False):
        # Get the directory where the executable is located
        exe_dir = os.path.dirname(sys.executable)
        debug_txt = os.path.join(exe_dir, "DEBUG.txt")
        debug_folder = os.path.join(exe_dir, "DEBUG")
        debug_prompts = os.path.join(debug_folder, "prompts.ini")
        
        # Check for DEBUG.txt file (only trigger for debug mode)
        if os.path.exists(debug_txt):
            debug_mode = True
            debug_print(f"DEBUG MODE: DEBUG.txt detected at {debug_txt}")
            
            # Create DEBUG folder if it doesn't exist
            os.makedirs(debug_folder, exist_ok=True)
            
            # Use prompts.ini from DEBUG folder
            if os.path.exists(debug_prompts):
                prompts_path = debug_prompts
                debug_print(f"DEBUG MODE: Using prompts.ini from DEBUG folder")
            else:
                # Copy built-in prompts to DEBUG folder
                import shutil
                shutil.copy2(_resource_path("prompts.ini"), debug_prompts)
                prompts_path = debug_prompts
                debug_print(f"DEBUG MODE: Copied built-in prompts.ini to DEBUG folder")
    
    # Use built-in prompts.ini if no debug mode
    if prompts_path is None:
        prompts_path = _resource_path("prompts.ini")
    
    return prompts_path, debug_mode

TEMPLATE_PATH  = _resource_path("B2 templateee.docx")
PROMPTS_INI, DEBUG_MODE = _get_prompts_path()
FEEDBACK_FILE  = _resource_path("feedback.ini")
LEVEL_TOKENS   = ["A1","A2","B1","B1+","B2","C1","C2"]

# Toggle debug prints for pixmap sizes/DPR
DEBUG_PREVIEW = False

def debug_print(*a):
    if DEBUG_PREVIEW:
        print(*a)

def debug_mode_info():
    """Print debug mode information"""
    if DEBUG_MODE:
        debug_print("=" * 50)
        debug_print("DEBUG MODE ACTIVATED")
        debug_print(f"Prompts source: {PROMPTS_INI}")
        debug_print("Using prompts.ini from DEBUG folder")
        debug_print("AI conversations will be logged to DEBUG folder.")
        debug_print("Debug mode triggered by: DEBUG.txt file next to executable")
        debug_print("=" * 50)

def setup_debug_folder():
    """Get DEBUG folder path for logging"""
    if not DEBUG_MODE:
        return None
    
    try:
        # Get the directory where the executable is located
        exe_dir = os.path.dirname(sys.executable)
        debug_dir = os.path.join(exe_dir, "DEBUG")
        return debug_dir
    except Exception as e:
        debug_print(f"DEBUG: Failed to get DEBUG folder path: {e}")
        return None

def log_ocr_output(file_path, file_type, ocr_text):
    """Log OCR output to debug file"""
    if not DEBUG_MODE:
        return
    
    try:
        debug_dir = setup_debug_folder()
        if not debug_dir:
            return
        
        # Create timestamped log file
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = os.path.join(debug_dir, f"ocr_output_{timestamp}.txt")
        
        # Format the log entry
        log_entry = f"""
{'='*80}
TIMESTAMP: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
OCR FILE: {file_path}
FILE TYPE: {file_type.upper()}
{'='*80}

EXTRACTED TEXT:
{ocr_text}

{'='*80}
END OF OCR OUTPUT
{'='*80}
"""
        
        # Write to log file
        with open(log_file, 'w', encoding='utf-8') as f:
            f.write(log_entry)
        
        debug_print(f"[DEBUG] OCR output logged to: {log_file}")
        
    except Exception as e:
        debug_print(f"[DEBUG] Failed to log OCR output: {e}")

def log_ai_conversation(conversation_type, prompt, response, error=None):
    """Log AI conversations to debug file"""
    if not DEBUG_MODE:
        return
    
    try:
        debug_dir = setup_debug_folder()
        if not debug_dir:
            return
        
        # Create timestamped log file
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = os.path.join(debug_dir, f"ai_conversations_{timestamp}.txt")
        
        # Check if prompt contains unformatted placeholders
        has_placeholders = any(placeholder in prompt for placeholder in ['{topics}', '{vocab}', '{unit_title}', '{unit_level}', '{cefr_tier}'])
        
        # Format the log entry
        log_entry = f"""
{'='*80}
TIMESTAMP: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
CONVERSATION TYPE: {conversation_type}
FORMATTING STATUS: {'WARNING - Contains unformatted placeholders!' if has_placeholders else 'OK - Properly formatted'}
{'='*80}

PROMPT:
{prompt}

{'='*40}
RESPONSE:
{response if response else 'No response'}

{'='*40}
ERROR:
{error if error else 'No error'}

{'='*80}

"""
        
        # Append to log file
        with open(log_file, "a", encoding="utf-8") as f:
            f.write(log_entry)
            
        # Also print warning to console if formatting failed
        if has_placeholders:
            debug_print(f"WARNING: Logged prompt contains unformatted placeholders!")
            debug_print(f"Check log file: {log_file}")
            
    except Exception as e:
        debug_print(f"DEBUG: Failed to log AI conversation: {e}")

# =========================
# Updater (called from About…)
# =========================
def version_tuple(v: str):
    nums = [int(x) for x in re.findall(r"\d+", v or "0")]
    return tuple(nums) or (0,)

def current_executable_path() -> str:
    return sys.executable if getattr(sys, "frozen", False) else os.path.abspath(sys.argv[0])

def fetch_manifest(url: str, timeout: int) -> dict:
    # Normalize GitHub blob → raw if needed
    if url and "github.com" in url and "/blob/" in url:
        url = url.replace("https://github.com/", "https://raw.githubusercontent.com/").replace("/blob/", "/")
    req = Request(url, headers={"User-Agent": f"{APP_NAME}/{APP_VERSION}", "Accept":"application/json"})
    with urlopen(req, timeout=timeout) as resp:
        data = resp.read()
    if b"<html" in data[:200].lower():
        raise RuntimeError("Invalid manifest (HTML received)")
    return json.loads(data.decode("utf-8"))

def sha256_of_file(path: str) -> str:
    import hashlib as _h
    h = _h.sha256()
    with open(path, "rb") as f:
        for b in iter(lambda: f.read(1024 * 1024), b""):
            h.update(b)
    return h.hexdigest()

class _Downloader(QThread):
    progress = pyqtSignal(int, int)  # read, total
    done     = pyqtSignal(str)       # path
    err      = pyqtSignal(str)
    def __init__(self, url: str):
        super().__init__()
        # normalize GH raw
        if url and "github.com" in url and "/blob/" in url:
            url = url.replace("https://github.com/", "https://raw.githubusercontent.com/").replace("/blob/", "/")
        self.url = url
        fd, p = tempfile.mkstemp(prefix="upd_", suffix=".bin"); os.close(fd)
        self.out = p
    def run(self):
        try:
            req = Request(self.url, headers={"User-Agent": f"{APP_NAME}/{APP_VERSION}", "Accept":"*/*"})
            with urlopen(req, timeout=NETWORK_TIMEOUT_SEC) as resp:
                total = resp.headers.get("Content-Length")
                total = int(total) if total and total.isdigit() else -1
                read = 0
                chunk = 1024 * 256
                with open(self.out, "wb") as f:
                    while True:
                        buf = resp.read(chunk)
                        if not buf: break
                        f.write(buf); read += len(buf)
                        self.progress.emit(read, total)
            self.done.emit(self.out)
        except Exception as e:
            self.err.emit(str(e))

class _ProgressDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("EFL Cafe Wizard - Downloading Update")
        self.setModal(True)
        self.setWindowFlags(Qt.Dialog | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
        lay = QVLayout(self)
        self.label = QLabel("Starting…")
        self.bar   = QProgressBar(); self.bar.setMinimum(0); self.bar.setMaximum(0)
        lay.addWidget(self.label); lay.addWidget(self.bar)
        self.resize(420,140)
    def set_progress(self, read, total):
        if total>0:
            pct = max(0, min(100, int(read*100/total)))
            self.bar.setMaximum(100); self.bar.setValue(pct)
            self.label.setText(f"Downloading… {pct}% ({read//1024//1024} MB / {total//1024//1024} MB)")
        else:
            self.bar.setMaximum(0); self.label.setText(f"Downloading… {read//1024//1024} MB")

class UpdateAvailableDialog(QDialog):
    """Custom dialog for showing update availability with rich text formatting"""
    def __init__(self, parent=None, latest_version="", current_version="", notes=""):
        super().__init__(parent)
        self.setWindowTitle("EFL Cafe Wizard - Update Available")
        self.setModal(True)
        self.setWindowFlags(Qt.Dialog | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
        self.resize(450, 250)
        
        layout = QVLayout(self)
        
        # Title
        title = QLabel("🔄 Update Available")
        title.setFont(QFont("Tahoma", 12, QFont.Bold))
        title.setStyleSheet("color: #007acc; margin-bottom: 10px;")
        layout.addWidget(title)
        
        # Version info with bold formatting
        version_text = f"New version: <b>{latest_version}</b><br>Current: <b>{current_version}</b>"
        version_label = QLabel(version_text)
        version_label.setTextFormat(Qt.RichText)
        version_label.setFont(QFont("Tahoma", 9))
        layout.addWidget(version_label)
        
        # Patch notes section
        if notes:
            notes_title = QLabel("📋 PATCH NOTES:")
            notes_title.setFont(QFont("Tahoma", 9, QFont.Bold))
            notes_title.setStyleSheet("color: #333; margin-top: 10px;")
            layout.addWidget(notes_title)
            
            # Format notes
            if isinstance(notes, list):
                notes_text = "<br>".join([f"• {note}" for note in notes])
            else:
                notes_text = notes.strip()
            
            notes_label = QLabel(notes_text)
            notes_label.setTextFormat(Qt.RichText)
            notes_label.setFont(QFont("Tahoma", 10))
            notes_label.setWordWrap(True)
            notes_label.setStyleSheet("background: #f5f5f5; padding: 8px; border-radius: 4px; margin: 5px 0;")
            layout.addWidget(notes_label)
        
        # Buttons
        button_layout = QHBoxLayout()
        self.btn_update = QPushButton("Update Now")
        self.btn_update.setStyleSheet(get_standard_button_style("dialog"))
        self.btn_later = QPushButton("Later")
        self.btn_later.setStyleSheet(get_standard_button_style("dialog_secondary"))
        
        self._update_in_progress = False
        self.btn_update.clicked.connect(self.start_update)
        self.btn_later.clicked.connect(self.reject)
        
        button_layout.addStretch()
        button_layout.addWidget(self.btn_later)
        button_layout.addWidget(self.btn_update)
        layout.addLayout(button_layout)
    
    def start_update(self):
        """Start the update process with proper state management"""
        if self._update_in_progress:
            return
        
        self._update_in_progress = True
        self.btn_update.setEnabled(False)
        self.btn_update.setText("Updating...")
        self.btn_later.setEnabled(False)
        
        # Show progress message
        QMessageBox.information(self, "Update", 
            "Starting update process...\n\n"
            "Please wait while the update is prepared.\n"
            "Do not close the application during this process.")
        
        # Accept the dialog to proceed with update
        self.accept()

def begin_update_replace(new_exe_path: str, target_exe: str) -> None:
    global UPDATE_IN_PROGRESS
    
    try:
        if os.name == "nt":
            # Use a more reliable update method without elevation
            new_path = target_exe + ".new"
            
            # Create a lock file to prevent multiple update attempts
            lock_file = target_exe + ".update_lock"
            try:
                # Check if another update is already in progress
                if os.path.exists(lock_file):
                    QMessageBox.warning(None, "Update", "Another update is already in progress. Please wait for it to complete.")
                    return
                
                # Create lock file
                with open(lock_file, "w") as f:
                    f.write(f"Update started at {__import__('datetime').datetime.now()}")
                
                # Copy the new executable with progress indication
                QMessageBox.information(None, "Update", 
                    f"Preparing update...\n\n"
                    f"This may take a few moments.\n"
                    f"Please do not close the application.")
                
                # Copy with retry mechanism
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        shutil.copyfile(new_exe_path, new_path)
                        break
                    except Exception as e:
                        if attempt == max_retries - 1:
                            raise e
                        import time
                        time.sleep(1)
                
                # Verify the copy was successful
                if not os.path.exists(new_path):
                    raise RuntimeError("Failed to create .new file")
                
                # Create a batch script for the update process (without auto-restart)
                bat_path = os.path.join(tempfile.gettempdir(), f"update_{os.getpid()}.bat")
                script = f"""@echo off
timeout /t 3 /nobreak >nul >nul 2>&1
:retry
if exist "{target_exe}" (
    del /f /q "{target_exe}" >nul 2>&1
    if errorlevel 1 (
        timeout /t 2 /nobreak >nul >nul 2>&1
        goto retry
    )
)
move "{new_path}" "{target_exe}" >nul 2>&1
if errorlevel 1 (
    timeout /t 2 /nobreak >nul >nul 2>&1
    goto retry
)
del /f /q "{lock_file}" >nul 2>&1
del /f /q "%~f0" >nul 2>&1
"""
                with open(bat_path, "w", encoding="utf-8") as f: f.write(script)
                
                # Run the batch script without showing a window
                creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
                subprocess.Popen([bat_path], creationflags=creationflags, close_fds=True, 
                               stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                
                QMessageBox.information(None, "Update Complete", 
                    f"Update completed successfully!\n\n"
                    f"Please manually restart {APP_NAME} to use the new version.\n\n"
                    f"The application will now close.")
                QApplication.quit(); sys.exit(0)
                
            finally:
                # Clean up lock file if something went wrong
                try:
                    if os.path.exists(lock_file):
                        os.remove(lock_file)
                except:
                    pass
        else:
            # Non-Windows: simpler approach
            shutil.copyfile(new_exe_path, target_exe + ".new")
            QMessageBox.information(None, "Update Complete", 
                f"Update completed successfully!\n\n"
                f"Please manually restart {APP_NAME} to use the new version.\n\n"
                f"The application will now close.")
            QApplication.quit(); sys.exit(0)
    except Exception as e:
        UPDATE_IN_PROGRESS = False
        QMessageBox.critical(None, "Updater", f"Failed to finalize update:\n{e}")

class BackgroundUpdateChecker(QThread):
    """Background thread for checking updates without blocking the UI"""
    update_available = pyqtSignal(dict)  # manifest data
    check_completed = pyqtSignal(bool)   # True if update available, False if up to date
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self._running = True
    
    def run(self):
        try:
            if not AUTO_UPDATE_ENABLED:
                return
                
            m = fetch_manifest(UPDATE_MANIFEST_URL, NETWORK_TIMEOUT_SEC)
            latest = m.get("version", "0")
            if version_tuple(latest) > version_tuple(APP_VERSION):
                self.update_available.emit(m)
                self.check_completed.emit(True)
            else:
                self.check_completed.emit(False)
        except Exception as e:
            debug_print(f"Background update check failed: {e}")
            self.check_completed.emit(False)
    
    def stop(self):
        self._running = False

def check_for_updates(parent, show_updater_pane=False):
    global UPDATE_IN_PROGRESS
    
    # Prevent multiple update dialogs
    if UPDATE_IN_PROGRESS:
        QMessageBox.information(parent, "Update", "An update is already in progress. Please wait for it to complete.")
        return False
    
    try:
        m = fetch_manifest(UPDATE_MANIFEST_URL, NETWORK_TIMEOUT_SEC)
        latest = m.get("version", "0"); url = m.get("url"); sha = m.get("sha256"); notes = m.get("notes")
        if version_tuple(latest) <= version_tuple(APP_VERSION):
            if not show_updater_pane:
                QMessageBox.information(parent, "Updater", f"You're on the latest version (v{APP_VERSION}).")
            return False
        
        # Use custom dialog with rich text formatting
        update_dialog = UpdateAvailableDialog(parent, latest, APP_VERSION, notes)
        if update_dialog.exec_() != QDialog.Accepted:
            return False
        
        # Set update in progress immediately after dialog acceptance
        UPDATE_IN_PROGRESS = True
            
        dlg = _ProgressDialog(parent); dl = _Downloader(url)
        dl.progress.connect(dlg.set_progress)
        def on_done(path):
            try:
                if sha:
                    got = sha256_of_file(path)
                    if got.lower() != sha.lower():
                        dlg.close(); QMessageBox.critical(parent, "Updater", "SHA-256 mismatch. Update aborted."); 
                        UPDATE_IN_PROGRESS = False; return
                dlg.close(); QMessageBox.information(parent, "Updater", "Update downloaded. The app will exit to finalize.")
                begin_update_replace(path, current_executable_path())
            finally:
                pass
        def on_err(msg): 
            dlg.close(); QMessageBox.critical(parent, "Download Failed", msg)
            UPDATE_IN_PROGRESS = False
        dl.done.connect(on_done); dl.err.connect(on_err)
        dl.start(); dlg.show()
        return True
    except Exception as e:
        UPDATE_IN_PROGRESS = False
        QMessageBox.critical(parent, "Updater", str(e))
        return False

# =========================
# Prompts & Parsing
# =========================
DEFAULT_PROMPTS = {
    "batch": (
        "You are an ESL materials writer for an online {modifier}{unit_level} discussion class.\n"
        "UNIT: {unit_title}\nCONTENT TOPICS:\n{topics}\nTARGET VOCABULARY (optional): {vocab}\n"
        "CEFR TIER: {cefr_tier}\n\n"
        "RULES\n- EXACTLY 15 questions, numbered 1–15, one per line.\n"
        "- Each is ONE sentence, 12–20 words, ends with '?'.\n"
        "- Use the topics; distribute evenly; high-frequency English, CEFR {unit_level}.\n"
        "- {tier_instructions}\n"
        "- Avoid clichés/filler (amazing/awesome/etc.).\n"
        "OUTPUT: only the 15 numbered questions."
    ),
    "single": (
        "You are an ESL materials writer for an online {modifier}{unit_level} discussion class.\n"
        "UNIT: {unit_title}\nCONTENT TOPICS:\n{topics}\nTARGET VOCABULARY (optional): {vocab}\n"
        "CEFR TIER: {cefr_tier}\n\n"
        "TASK: Write ONE NEW question (12–20 words, one sentence, ends with '?').\n"
        "- {tier_instructions}\n"
        "Avoid clichés. Do not repeat any of these:\n{existing_questions}\n"
        "OUTPUT: only the question text."
    ),
}

from configparser import ConfigParser
def load_prompts(path):
    cfg = ConfigParser()
    if os.path.exists(path):
        try:
            cfg.read(path, encoding="utf-8")
        except Exception:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                cfg.read_file(f)
    
    # Check version compatibility in debug mode
    if DEBUG_MODE:
        check_prompts_version(path)
    
    def get(sec, default): return cfg.get(sec, "template", fallback=default) if cfg.has_section(sec) else default
    return {"batch": get("batch", DEFAULT_PROMPTS["batch"]),
            "single": get("single", DEFAULT_PROMPTS["single"])}

def check_prompts_version(prompts_path):
    """Check if prompts.ini version matches the application version in debug mode"""
    if not os.path.exists(prompts_path):
        debug_print(f"[DEBUG] WARNING: prompts.ini not found at {prompts_path}")
        return
    
    try:
        with open(prompts_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        # Extract version from the file header
        version_match = re.search(r'# Version:\s*([^\n]+)', content)
        if version_match:
            file_version = version_match.group(1).strip()
            if file_version != APP_VERSION:
                debug_print(f"[DEBUG] ⚠️  VERSION MISMATCH WARNING:")
                debug_print(f"[DEBUG]    Application Version: {APP_VERSION}")
                debug_print(f"[DEBUG]    Prompts File Version: {file_version}")
                debug_print(f"[DEBUG]    File Path: {prompts_path}")
                debug_print(f"[DEBUG]    Please update your prompts.ini file to match the application version.")
                debug_print(f"[DEBUG]    This may cause unexpected behavior or errors.")
                
                # Offer to auto-update the version
                try:
                    update_prompts_version(prompts_path, content)
                except Exception as e:
                    debug_print(f"[DEBUG]    Failed to auto-update version: {e}")
            else:
                debug_print(f"[DEBUG] ✅ Prompts version check passed: {file_version}")
        else:
            debug_print(f"[DEBUG] ⚠️  WARNING: No version information found in prompts.ini")
            debug_print(f"[DEBUG]    Expected version: {APP_VERSION}")
            debug_print(f"[DEBUG]    File: {prompts_path}")
    except Exception as e:
        debug_print(f"[DEBUG] ⚠️  Error checking prompts version: {e}")

def get_prompts_version():
    """Get the version from prompts.ini file"""
    try:
        if os.path.exists(PROMPTS_INI):
            with open(PROMPTS_INI, "r", encoding="utf-8") as f:
                content = f.read()
            
            version_match = re.search(r'# Version:\s*([^\n]+)', content)
            if version_match:
                return version_match.group(1).strip()
    except Exception:
        pass
    return None

def update_prompts_version(prompts_path, content):
    """Auto-update the version in prompts.ini to match the application version"""
    try:
        # Create a backup of the original file
        backup_path = prompts_path + ".backup"
        with open(backup_path, "w", encoding="utf-8") as f:
            f.write(content)
        debug_print(f"[DEBUG] 📁 Created backup: {backup_path}")
        
        # Update the version line
        updated_content = re.sub(
            r'# Version:\s*[^\n]+',
            f'# Version: {APP_VERSION}',
            content
        )
        
        # Update the last updated date
        from datetime import datetime
        current_date = datetime.now().strftime("%Y-%m-%d")
        updated_content = re.sub(
            r'# Last Updated:\s*[^\n]+',
            f'# Last Updated: {current_date}',
            updated_content
        )
        
        # Write the updated content back to the file
        with open(prompts_path, "w", encoding="utf-8") as f:
            f.write(updated_content)
        
        debug_print(f"[DEBUG] ✅ Auto-updated prompts.ini version to {APP_VERSION}")
        debug_print(f"[DEBUG]    Backup saved as: {backup_path}")
        
    except Exception as e:
        debug_print(f"[DEBUG] ❌ Failed to auto-update prompts.ini: {e}")
        raise

def fmt(t, **kw):
    try: 
        result = t.format(**kw)
        return result
    except Exception as e:
        debug_print(f"ERROR: Prompt formatting failed: {e}")
        debug_print(f"Template: {t[:200]}...")
        debug_print(f"Available keys: {list(kw.keys())}")
        return t

def get_tier_instructions(level, tier):
    """Generate specific instructions based on CEFR level and tier"""
    if tier == "Lower":
        if level in ["A1", "A2"]:
            return "Use simple present tense, basic vocabulary, and straightforward questions like 'Do you...?', 'Can you...?', 'Do you like...?'"
        elif level in ["B1", "B1+"]:
            return "Use simple to intermediate structures, common vocabulary, and questions like 'Do you think...?', 'Have you ever...?', 'Would you like to...?'"
        else:  # B2, C1, C2
            return "Use intermediate structures, avoid overly complex grammar, and focus on practical, everyday questions"
    elif tier == "Neutral":
        if level in ["A1", "A2"]:
            return "Use appropriate structures for the level, mix of present and past tense, and balanced questions that are neither too simple nor too complex"
        elif level in ["B1", "B1+"]:
            return "Use intermediate structures, standard vocabulary, and well-balanced questions that match the level expectations"
        else:  # B2, C1, C2
            return "Use level-appropriate structures, standard vocabulary, and questions that are challenging but not overly complex"
    else:  # Upper
        if level in ["A1", "A2"]:
            return "Use more varied structures within the level, include past tense, and create engaging questions beyond basic patterns"
        elif level in ["B1", "B1+"]:
            return "Use intermediate to upper-intermediate structures, more sophisticated vocabulary, and thought-provoking questions"
        else:  # B2, C1, C2
            return "Use advanced structures, complex vocabulary, and challenging questions that require critical thinking and detailed responses"

def get_quality_validation_instructions(enabled):
    """Generate quality validation instructions"""
    if not enabled:
        return ""
    return """QUALITY VALIDATION REQUIREMENTS
	- Ensure each question is crystal clear and unambiguous
	- Use perfect grammar and natural phrasing throughout
	- Avoid culturally insensitive or potentially offensive content
	- Make sure questions are appropriate for the target level
	- Double-check that questions sound like natural English"""

def get_blooms_taxonomy_instructions(level, blooms_level):
    """Generate Bloom's taxonomy instructions"""
    if blooms_level == "Auto":
        return ""
    
    instructions = f"BLOOM'S TAXONOMY LEVEL: {blooms_level.upper()}\n"
    
    if blooms_level == "Remember":
        instructions += """	- Focus on factual recall questions
	- Use stems like: What is...?, Who...?, When...?, Where...?, Which...?
	- Ask for basic information and definitions"""
    elif blooms_level == "Understand":
        instructions += """	- Focus on comprehension and explanation
	- Use stems like: Explain why...?, Describe...?, What does... mean?
	- Ask students to interpret or summarize information"""
    elif blooms_level == "Apply":
        instructions += """	- Focus on practical application
	- Use stems like: How would you use...?, Solve...?, What would happen if...?
	- Ask students to use knowledge in new situations"""
    elif blooms_level == "Analyze":
        instructions += """	- Focus on breaking down and comparing
	- Use stems like: Compare...?, What are the differences...?, Why do you think...?
	- Ask students to examine relationships and patterns"""
    elif blooms_level == "Evaluate":
        instructions += """	- Focus on judgment and opinion
	- Use stems like: Do you agree...?, What's your opinion...?, Which is better...?
	- Ask students to make judgments and defend positions"""
    elif blooms_level == "Create":
        instructions += """	- Focus on original thinking and design
	- Use stems like: Design...?, Invent...?, Propose...?, What would you create...?
	- Ask students to generate new ideas or solutions"""
    
    return instructions

def get_engagement_level_instructions(engagement_level):
    """Generate engagement level instructions"""
    if engagement_level == "Balanced":
        return ""
    
    instructions = f"ENGAGEMENT LEVEL: {engagement_level.upper()}\n"
    
    if engagement_level == "Low":
        instructions += """	- Create simple, factual questions
	- Focus on basic information and personal preferences
	- Avoid controversial or complex topics"""
    elif engagement_level == "Medium":
        instructions += """	- Include personal but safe questions
	- Mix factual and opinion-based questions
	- Use topics that most students can relate to"""
    elif engagement_level == "High":
        instructions += """	- Create thought-provoking, opinion-based questions
	- Include questions that spark debate and discussion
	- Use topics that require deeper thinking"""
    elif engagement_level == "Very High":
        instructions += """	- Create controversial and challenging questions
	- Include topics that require strong opinions and debate
	- Use complex, multi-faceted issues that generate intense discussion"""
    
    return instructions

def get_academic_background_instructions(enabled):
    """Generate academic background instructions"""
    if not enabled:
        return ""
    return """ACADEMIC BACKGROUND REQUIREMENT
	- Generate questions that require subject-specific knowledge
	- Use academic vocabulary and terminology appropriately
	- Include questions that test critical thinking skills
	- Require students to draw on research or study background
	- Make questions suitable for advanced or specialized learners"""

def get_naturalness_instructions(enabled):
    """Generate naturalness check instructions"""
    if not enabled:
        return ""
    return """NATURALNESS REQUIREMENTS
	- Use native speaker language patterns and phrasing
	- Ensure appropriate register for the target level
	- Create natural question flow and structure
	- Avoid awkward, artificial, or textbook-like phrasing
	- Make questions sound like they come from a real teacher"""

def get_strictness_instructions(strictness):
    """Generate topic consistency instructions based on strictness level"""
    if strictness == "Strict":
        return """TOPIC CONSISTENCY: STRICT
	- Questions must strictly follow the provided topics
	- No deviation from topic content
	- Questions must directly relate to specific topic points
	- Minimal creative interpretation allowed"""
    elif strictness == "Free":
        return """TOPIC CONSISTENCY: FREE
	- More flexible interpretation of topics
	- Questions can be loosely related to topics
	- Allows creative and tangential connections
	- Good for open-ended discussions"""
    else:  # Neutral (default)
        return """TOPIC CONSISTENCY: NEUTRAL
	- Questions should relate to topics but allow some flexibility
	- Good balance between strict adherence and creativity
	- Recommended for most use cases"""

def parse_numbered_questions(text, max_items=15):
    out = []
    for line in (text or "").splitlines():
        s = line.strip()
        m = re.match(r'^\s*(\d+)[\.\)]\s+(.*)\?$', s)
        if m: out.append(m.group(2).strip() + "?"); continue
        m2 = re.match(r'^[\-\*\u2022]\s+(.*)\?$', s)
        if m2: out.append(m2.group(1).strip() + "?")
    if not out:
        parts = re.split(r'\s*\d+[\.\)]\s+', text or "")
        for c in parts:
            c = c.strip()
            if c.endswith("?"): out.append(c)
    clean, seen = [], set()
    for q in out[:max_items]:
        q = re.sub(r"\s+", " ", q).strip().strip('"”').rstrip()
        if not q.endswith("?"): q = q.rstrip(".! ") + "?"
        if q.lower() in seen: continue
        seen.add(q.lower()); clean.append(q)
    return clean

SPEC_VOCAB_RE = re.compile(r'(?i)^vocab\s*:\s*(.*)$')
UNIT_LINE_RE  = re.compile(r'(?i)^\s*unit\s*\d+[A-B]?\s*[-–]\s*(.+)$')

def parse_topics_spec(raw_text):
    unit_title = None
    topics, vocab = [], []
    lines = [l.rstrip() for l in (raw_text or "").splitlines()]
    i = 0; vocab_mode = False; vocab_chunks = []
    while i < len(lines):
        line = lines[i].strip()
        if not line: i += 1; continue
        if unit_title is None:
            m = UNIT_LINE_RE.match(line)
            if m: unit_title = m.group(1).strip(); i += 1; continue
            if line.lower().startswith("unit") and "-" in line:
                try: unit_title = line.split("-", 1)[1].strip(); i += 1; continue
                except Exception: pass
        m = SPEC_VOCAB_RE.match(line)
        if m:
            first = m.group(1).strip()
            if first: vocab_chunks.append(first)
            vocab_mode = True; i += 1; continue
        if vocab_mode:
            if line.startswith(("*","-")) or line.lower().startswith("unit"):
                vocab_mode = False
            else:
                vocab_chunks.append(line); i += 1; continue
        if line.startswith(("*","-")):
            topics.append(line.lstrip("*- ").strip())
        elif not line.lower().startswith("vocab"):
            if len(line.split()) <= 12 and not line.endswith("?"):
                topics.append(line.strip())
        i += 1
    if vocab_chunks:
        alltext = " ".join(vocab_chunks)
        vocab = [w.strip(" .;:,") for w in re.split(r'[,\"\"\u2022;]+', alltext) if w.strip()]
        dedup, seen = [], set()
        for w in vocab:
            lw = w.lower()
            if lw in seen: continue
            seen.add(lw); dedup.append(w)
        vocab = dedup
    topics = [re.sub(r"\s+", " ", t).strip(" .") for t in topics if t.strip()]
    topics = [t for t in topics if not t.lower().startswith("vocab")]
    return (unit_title if unit_title else None, topics, vocab)

# =========================
# OCR functionality
# =========================

def setup_tesseract_path():
    """Setup Tesseract path for Windows if needed."""
    if not OCR_AVAILABLE:
        return False
    
    # Check if tesseract is already in PATH
    try:
        pytesseract.get_tesseract_version()
        return True
    except:
        pass
    
    # Get the base path for the application
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        base_path = sys._MEIPASS
    else:
        # Running as script
        base_path = os.path.dirname(os.path.abspath(__file__))
    
    # Try paths in order of preference
    possible_paths = [
        # PyInstaller bundled path (highest priority)
        os.path.join(base_path, "Tesseract-OCR", "tesseract.exe"),
        # Local development path
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "Tesseract-OCR", "tesseract.exe"),
        # Relative path from current directory
        r".\Tesseract-OCR\tesseract.exe",
        # System installation paths
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\Users\{}\AppData\Local\Tesseract-OCR\tesseract.exe".format(os.getenv('USERNAME', ''))
    ]
    
    for path in possible_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            try:
                pytesseract.get_tesseract_version()
                debug_print(f"[OCR] Using Tesseract at: {path}")
                return True
            except Exception as e:
                debug_print(f"[OCR] Failed to use Tesseract at {path}: {e}")
                continue
    
    debug_print("[OCR] Could not find Tesseract executable")
    return False

def setup_tessdata_path():
    """Setup tessdata path for compiled environment."""
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        base_path = sys._MEIPASS
        tessdata_path = os.path.join(base_path, "Tesseract-OCR", "tessdata")
    else:
        # Running as script
        tessdata_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Tesseract-OCR", "tessdata")
    
    if os.path.exists(tessdata_path):
        # Set the tessdata path for pytesseract
        os.environ['TESSDATA_PREFIX'] = tessdata_path
        debug_print(f"[OCR] Using tessdata at: {tessdata_path}")
        return True
    else:
        debug_print(f"[OCR] Tessdata not found at: {tessdata_path}")
        return False

def ocr_image_to_text(image_path, dpi=200, lang="eng", config="--psm 3"):
    """Extract text from image using OCR."""
    if not OCR_AVAILABLE:
        raise ImportError("OCR dependencies not available")
    
    if not setup_tesseract_path():
        raise RuntimeError("Tesseract not found. Please install Tesseract OCR.")
    
    # Setup tessdata path for compiled environment
    setup_tessdata_path()
    
    try:
        # Open and process image
        with Image.open(image_path) as img:
            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Apply grayscale for better OCR
            img = ImageOps.grayscale(img)
            
            # Perform OCR with hidden console window
            if sys.platform.startswith('win'):
                # Windows: Use subprocess with CREATE_NO_WINDOW flag to hide console
                import subprocess
                import tempfile
                
                # Save image to temporary file
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    img.save(tmp_file.name)
                    tmp_path = tmp_file.name
                
                try:
                    # Run tesseract with hidden console
                    cmd = [
                        pytesseract.pytesseract.tesseract_cmd,
                        tmp_path,
                        'stdout',
                        '-l', lang,
                        '--psm', '3'
                    ]
                    
                    # Use CREATE_NO_WINDOW to prevent console flash
                    result = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
                    )
                    
                    if result.returncode == 0:
                        text = result.stdout.strip()
                    else:
                        raise RuntimeError(f"Tesseract failed: {result.stderr}")
                        
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(tmp_path)
                    except:
                        pass
            else:
                # Non-Windows: Use regular pytesseract
                text = pytesseract.image_to_string(img, lang=lang, config=config)
            
            return text.strip()
    except Exception as e:
        raise RuntimeError(f"OCR failed: {str(e)}")

def ocr_pdf_to_text(pdf_path, dpi=150, lang="eng", config="--psm 3", max_pages=None):
    """Extract text from PDF using OCR."""
    if not OCR_AVAILABLE:
        raise ImportError("OCR dependencies not available")
    
    if not setup_tesseract_path():
        raise RuntimeError("Tesseract not found. Please install Tesseract OCR.")
    
    # Setup tessdata path for compiled environment
    setup_tessdata_path()
    
    try:
        doc = fitz.open(pdf_path)
        texts = []
        
        # Limit pages if specified
        page_count = len(doc)
        if max_pages:
            page_count = min(page_count, max_pages)
        
        for i in range(page_count):
            page = doc[i]
            # Render page to image
            pix = page.get_pixmap(dpi=dpi)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            
            # Apply grayscale for better OCR
            img = ImageOps.grayscale(img)
            
            # Perform OCR with hidden console window
            if sys.platform.startswith('win'):
                # Windows: Use subprocess with CREATE_NO_WINDOW flag to hide console
                import subprocess
                import tempfile
                
                # Save image to temporary file
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    img.save(tmp_file.name)
                    tmp_path = tmp_file.name
                
                try:
                    # Run tesseract with hidden console
                    cmd = [
                        pytesseract.pytesseract.tesseract_cmd,
                        tmp_path,
                        'stdout',
                        '-l', lang,
                        '--psm', '3'
                    ]
                    
                    # Use CREATE_NO_WINDOW to prevent console flash
                    result = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
                    )
                    
                    if result.returncode == 0:
                        text = result.stdout.strip()
                    else:
                        raise RuntimeError(f"Tesseract failed: {result.stderr}")
                        
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(tmp_path)
                    except:
                        pass
            else:
                # Non-Windows: Use regular pytesseract
                text = pytesseract.image_to_string(img, lang=lang, config=config)
            
            texts.append(f"=== PAGE {i+1} ===\n{text.strip()}")
        
        doc.close()
        return "\n\n".join(texts)
    except Exception as e:
        raise RuntimeError(f"PDF OCR failed: {str(e)}")

def extract_topics_from_text(text, unit_title="Document"):
    """Extract topics from OCR text using AI."""
    if not text or not text.strip():
        debug_print(f"[AI] No text provided for topic extraction")
        return unit_title, "", []
    
    # Use Gemini to extract topics from the OCR text
    try:
        if not GOOGLE_AI_AVAILABLE:
            debug_print(f"[AI] Google AI library not available, using HTTP fallback")
            return extract_topics_http_fallback(text, unit_title)
        
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        prompt = f"""Analyze the following text and extract simple discussion topics for ESL students.

1. A simple unit title (if not obvious, use "{unit_title}")
2. 3-4 simple discussion topics (NOT questions - just topic themes like "favorite foods" or "travel experiences")
3. 5-8 key vocabulary words (comma-separated)

IMPORTANT: 
- Topics should be THEMES for discussion, NOT questions
- Keep topics simple and conversational
- Examples: "favorite foods", "travel experiences", "daily routines", "hobbies"
- Avoid complex academic topics

Text to analyze:
{text[:2000]}

Format your response as:
UNIT_TITLE: [simple title]
TOPICS:
* [topic theme 1]
* [topic theme 2] 
* [topic theme 3]
VOCAB: [word1, word2, word3, ...]"""
        
        response = model.generate_content(prompt)
        
        # Handle different response formats more robustly
        result = None
        if hasattr(response, 'text') and response.text:
            result = response.text
        elif hasattr(response, 'candidates') and response.candidates:
            if len(response.candidates) > 0:
                candidate = response.candidates[0]
                if hasattr(candidate, 'content') and candidate.content:
                    if hasattr(candidate.content, 'parts') and candidate.content.parts:
                        if len(candidate.content.parts) > 0:
                            part = candidate.content.parts[0]
                            if hasattr(part, 'text') and part.text:
                                result = part.text
        
        if not result or not result.strip():
            debug_print(f"[AI] Empty or invalid response from Gemini API")
            debug_print(f"[AI] Response object: {response}")
            debug_print(f"[AI] Response type: {type(response)}")
            debug_print(f"[AI] Response attributes: {dir(response) if hasattr(response, '__dict__') else 'No attributes'}")
            return extract_topics_http_fallback(text, unit_title)
        
        result = result.strip()
        debug_print(f"[AI] Successfully extracted response: {len(result)} characters")
        
        # Parse the response
        if result is None:
            debug_print(f"[AI] Result is None, falling back to HTTP API")
            return extract_topics_http_fallback(text, unit_title)
        
        try:
            lines = result.split('\n')
        except AttributeError as e:
            debug_print(f"[AI] AttributeError when splitting lines: {e}")
            debug_print(f"[AI] Result type: {type(result)}, value: {result}")
            return extract_topics_http_fallback(text, unit_title)
        extracted_title = unit_title
        topics = []
        vocab = []
        
        for line in lines:
            line = line.strip()
            if line.startswith('UNIT_TITLE:'):
                extracted_title = line.replace('UNIT_TITLE:', '').strip()
            elif line.startswith('*'):
                topics.append(line[1:].strip())
            elif line.startswith('VOCAB:'):
                vocab_text = line.replace('VOCAB:', '').strip()
                vocab = [v.strip() for v in vocab_text.split(',') if v.strip()]
        
        # Format topics for the UI
        topics_text = ""
        for topic in topics:
            topics_text += f"* {topic}\n"
        
        return extracted_title, topics_text.strip(), vocab
        
    except AttributeError as e:
        if "'NoneType' object has no attribute 'splitlines'" in str(e):
            debug_print(f"[AI] NoneType splitlines error detected - Google AI response is None")
            debug_print(f"[AI] This is a common issue with compiled executables")
            debug_print(f"[AI] Falling back to HTTP-based AI")
            return extract_topics_http_fallback(text, unit_title)
        else:
            debug_print(f"[AI] AttributeError: {e}")
            debug_print(f"[AI] Falling back to HTTP-based AI")
            return extract_topics_http_fallback(text, unit_title)
    except Exception as e:
        debug_print(f"[AI] Google AI SDK failed: {e}")
        import traceback
        debug_print(f"[AI] Traceback: {traceback.format_exc()}")
        # Fallback: try HTTP-based AI first, then simple text analysis
        debug_print(f"[AI] Trying HTTP-based AI fallback")
        return extract_topics_http_fallback(text, unit_title)

def extract_topics_http_fallback(text, unit_title="Document"):
    """HTTP-based fallback for compiled executables using direct API calls."""
    debug_print(f"[AI] Using HTTP-based AI fallback for compiled executable")
    
    try:
        import urllib.request
        import json
        import ssl
        
        # Prepare the request
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent"
        
        headers = {
            'Content-Type': 'application/json',
            'x-goog-api-key': GEMINI_API_KEY
        }
        
        debug_print(f"[AI] Making HTTP request to Gemini API...")
        
        prompt = f"""Analyze the following text and extract simple discussion topics for ESL students.

1. A simple unit title (if not obvious, use "{unit_title}")
2. 3-4 simple discussion topics (NOT questions - just topic themes like "favorite foods" or "travel experiences")
3. 5-8 key vocabulary words (comma-separated)

IMPORTANT: 
- Topics should be THEMES for discussion, NOT questions
- Keep topics simple and conversational
- Examples: "favorite foods", "travel experiences", "daily routines", "hobbies"
- Avoid complex academic topics

Text to analyze:
{text[:2000]}

Format your response as:
UNIT_TITLE: [simple title]
TOPICS:
* [topic theme 1]
* [topic theme 2] 
* [topic theme 3]
VOCAB: [word1, word2, word3, ...]"""
        
        data = {
            "contents": [{
                "parts": [{"text": prompt}]
            }]
        }
        
        # Make the request with SSL context
        req = urllib.request.Request(url, data=json.dumps(data).encode('utf-8'), headers=headers)
        
        # Create SSL context for HTTPS
        ssl_context = ssl.create_default_context()
        
        with urllib.request.urlopen(req, timeout=30, context=ssl_context) as response:
            if response.status == 200:
                result = json.loads(response.read().decode('utf-8'))
                debug_print(f"[AI] HTTP request successful (status: {response.status})")
            else:
                debug_print(f"[AI] HTTP request failed with status: {response.status}")
                return extract_topics_simple_fallback(text, unit_title)
            
            if 'candidates' in result and len(result['candidates']) > 0:
                response_text = result['candidates'][0]['content']['parts'][0]['text']
                debug_print(f"[AI] HTTP API response received: {len(response_text)} characters")
                
                # Parse the response
                lines = response_text.split('\n')
                extracted_title = unit_title
                topics = []
                vocab = []
                
                for line in lines:
                    line = line.strip()
                    if line.startswith('UNIT_TITLE:'):
                        extracted_title = line.replace('UNIT_TITLE:', '').strip()
                    elif line.startswith('*'):
                        topics.append(line[1:].strip())
                    elif line.startswith('VOCAB:'):
                        vocab_text = line.replace('VOCAB:', '').strip()
                        vocab = [v.strip() for v in vocab_text.split(',') if v.strip()]
                
                # Format topics for the UI
                topics_text = ""
                for topic in topics:
                    topics_text += f"* {topic}\n"
                
                debug_print(f"[AI] HTTP fallback extracted - Title: '{extracted_title}', Topics: {len(topics)}, Vocab: {len(vocab)}")
                return extracted_title, topics_text.strip(), vocab
            else:
                debug_print(f"[AI] HTTP API returned unexpected response format")
                return extract_topics_simple_fallback(text, unit_title)
    
    except Exception as e:
        debug_print(f"[AI] HTTP fallback failed: {e}")
        return extract_topics_simple_fallback(text, unit_title)

def extract_topics_simple_fallback(text, unit_title="Document"):
    """Simple fallback topic extraction when AI is not available."""
    debug_print(f"[AI] Using simple text analysis fallback")
    
    # Generate simple, conversational topics
    words = text.lower().split()
    
    # Simple topic detection based on common words
    if any(word in words for word in ['food', 'eat', 'restaurant', 'cooking', 'meal']):
        topics = ["Favorite foods", "Cooking and recipes", "Restaurant experiences"]
    elif any(word in words for word in ['travel', 'trip', 'vacation', 'visit', 'journey']):
        topics = ["Travel experiences", "Dream destinations", "Travel planning"]
    elif any(word in words for word in ['work', 'job', 'career', 'office', 'business']):
        topics = ["Work and jobs", "Career goals", "Work-life balance"]
    elif any(word in words for word in ['family', 'parent', 'child', 'mother', 'father']):
        topics = ["Family relationships", "Family traditions", "Childhood memories"]
    elif any(word in words for word in ['sport', 'music', 'movie', 'book', 'game']):
        topics = ["Hobbies and interests", "Entertainment preferences", "Free time activities"]
    else:
        topics = ["Personal experiences", "Daily life", "General discussion"]
    
    # Extract some words as vocabulary (simple approach)
    words = text.lower().split()
    # Filter out common words and get unique words
    common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them'}
    vocab_words = []
    for word in words:
        word = word.strip('.,!?;:"()[]{}')
        if len(word) > 3 and word not in common_words and word.isalpha():
            vocab_words.append(word)
            if len(vocab_words) >= 8:  # Limit to 8 words
                break
    
    # Format topics for the UI
    topics_text = ""
    for topic in topics:
        topics_text += f"* {topic}\n"
    
    debug_print(f"[AI] Simple fallback extracted {len(topics)} topics and {len(vocab_words)} vocabulary words")
    
    return unit_title, topics_text.strip(), vocab_words

class OCRWorker(QThread):
    """Background worker for OCR processing."""
    finished = pyqtSignal(str, str, list)  # unit_title, topics_text, vocab_list
    error = pyqtSignal(str)
    progress = pyqtSignal(str)
    
    def __init__(self, file_path, file_type="image"):
        super().__init__()
        self.file_path = file_path
        self.file_type = file_type  # "image" or "pdf"
    
    def run(self):
        try:
            self.progress.emit("Starting OCR...")
            
            if not OCR_AVAILABLE:
                self.error.emit("OCR dependencies not available. Please install PIL, pytesseract, and PyMuPDF.")
                return
            
            # Debug information
            debug_print(f"[OCR] Processing {self.file_type}: {self.file_path}")
            debug_print(f"[OCR] Running as compiled: {getattr(sys, 'frozen', False)}")
            if getattr(sys, 'frozen', False):
                debug_print(f"[OCR] MEIPASS: {sys._MEIPASS}")
            
            # Perform OCR based on file type
            if self.file_type == "image":
                self.progress.emit("Processing image...")
                text = ocr_image_to_text(self.file_path)
            else:  # PDF
                self.progress.emit("Processing PDF...")
                text = ocr_pdf_to_text(self.file_path, max_pages=5)  # Limit to first 5 pages
            
            if not text.strip():
                self.error.emit("No text could be extracted from the document.")
                return
            
            debug_print(f"[OCR] Extracted text length: {len(text)} characters")
            
            # Log OCR output in debug mode
            if DEBUG_MODE:
                log_ocr_output(self.file_path, self.file_type, text)
            
            self.progress.emit("Extracting topics...")
            
            # Extract topics using AI
            debug_print(f"[AI] Starting topic extraction from {len(text)} characters of text")
            unit_title, topics_text, vocab_list = extract_topics_from_text(text)
            
            debug_print(f"[AI] Extracted - Title: '{unit_title}', Topics: {len(topics_text.split('*'))-1}, Vocab: {len(vocab_list)}")
            
            self.progress.emit("OCR completed successfully!")
            self.finished.emit(unit_title, topics_text, vocab_list)
            
        except Exception as e:
            debug_print(f"[OCR] Error: {str(e)}")
            import traceback
            debug_print(f"[OCR] Traceback: {traceback.format_exc()}")
            self.error.emit(f"OCR failed: {str(e)}")

# =========================
# Gemini workers
# =========================
def init_gemini():
    if not GEMINI_API_KEY or GEMINI_API_KEY == "YOUR_GEMINI_API_KEY":
        raise RuntimeError("Gemini API key is missing. Set GEMINI_API_KEY env var or edit GEMINI_API_KEY.")
    
    if not GOOGLE_AI_AVAILABLE:
        raise RuntimeError("Google AI library is not available in compiled executable. Use HTTP fallback.")
    
    genai.configure(api_key=GEMINI_API_KEY.strip())
    return genai.GenerativeModel(GEMINI_MODEL)

class GeminiBatchWorker(QThread):
    ok = pyqtSignal(list)      # questions list
    err = pyqtSignal(str)
    def __init__(self, unit_title, level, topics_text, vocab_list, cefr_tier="Upper", existing=None, modifier="", prompt_style="Standard", quality_validation=True, blooms_level="Auto", engagement_level="Balanced", academic_background=False, naturalness_check=True, topic_consistency="Neutral", prompts=None):
        super().__init__()
        self.unit_title = unit_title
        self.level = level
        self.topics_text = topics_text
        self.vocab_list = vocab_list or []
        self.cefr_tier = cefr_tier
        self.existing = existing or []
        self.modifier = modifier or ""
        self.prompt_style = prompt_style or "Standard"
        self.quality_validation = quality_validation
        self.blooms_level = blooms_level
        self.engagement_level = engagement_level
        self.academic_background = academic_background
        self.naturalness_check = naturalness_check
        self.topic_consistency = topic_consistency
        self.prompts = prompts or load_prompts(PROMPTS_INI)
    def run(self):
        try:
            model = init_gemini()
            
            # Get the template and format it
            template = self.prompts.get("batch")
            prompt = fmt(
                template,
                unit_level=self.level,
                level=self.level,  # Add level parameter for prompts.ini template
                modifier=self.modifier,  # Use modifier from advanced features (Upper /Lower /empty)
                unit_title=self.unit_title,
                topics=self.topics_text,
                vocab=", ".join(self.vocab_list) if self.vocab_list else "(none)",
                cefr_tier=self.cefr_tier,
                tier_instructions=get_tier_instructions(self.level, self.cefr_tier),
                quality_validation_instructions=get_quality_validation_instructions(self.quality_validation),
                blooms_taxonomy_instructions=get_blooms_taxonomy_instructions(self.level, self.blooms_level),
                engagement_level_instructions=get_engagement_level_instructions(self.engagement_level),
                academic_background_instructions=get_academic_background_instructions(self.academic_background),
                naturalness_instructions=get_naturalness_instructions(self.naturalness_check),
                strictness_instructions=get_strictness_instructions(self.topic_consistency)
            )
            
            # Check if formatting worked correctly
            if prompt == template:
                debug_print(f"WARNING: Prompt formatting may have failed. Template and result are identical.")
                debug_print(f"Template: {template[:200]}...")
                debug_print(f"Formatted: {prompt[:200]}...")
            
            resp = model.generate_content(prompt)
            text = getattr(resp, "text", None)
            if not text and hasattr(resp, "candidates") and resp.candidates and len(resp.candidates) > 0:
                candidate = resp.candidates[0]
                if hasattr(candidate, "content") and candidate.content and hasattr(candidate.content, "parts") and candidate.content.parts and len(candidate.content.parts) > 0:
                    text = candidate.content.parts[0].text
            qs = parse_numbered_questions(text or "", max_items=15)
            while len(qs) < 15:
                qs.append("How could this topic affect an everyday situation at home, work, or school?")
            
            # Log the conversation for debug mode
            log_ai_conversation(
                conversation_type="BATCH_GENERATION",
                prompt=prompt,
                response=text,
                error=None
            )
            
            self.ok.emit(qs[:15])
        except Exception as e:
            # Log the error for debug mode
            log_ai_conversation(
                conversation_type="BATCH_GENERATION",
                prompt=prompt if 'prompt' in locals() else "Failed to generate prompt",
                response=None,
                error=str(e)
            )
            self.err.emit(f"Gemini batch failed: {e}")

class GeminiSingleWorker(QThread):
    ok = pyqtSignal(int, str)  # index, question
    err = pyqtSignal(str)
    def __init__(self, index, unit_title, level, topics_text, vocab_list, existing, cefr_tier="Upper", feedback="", modifier="", prompt_style="Standard", quality_validation=True, blooms_level="Auto", engagement_level="Balanced", academic_background=False, naturalness_check=True, topic_consistency="Neutral", prompts=None):
        super().__init__()
        self.index = index
        self.unit_title = unit_title
        self.level = level
        self.topics_text = topics_text
        self.vocab_list = vocab_list or []
        self.cefr_tier = cefr_tier
        self.existing = existing or []
        self.feedback = feedback or ""
        self.modifier = modifier or ""
        self.prompt_style = prompt_style or "Standard"
        self.quality_validation = quality_validation
        self.blooms_level = blooms_level
        self.engagement_level = engagement_level
        self.academic_background = academic_background
        self.naturalness_check = naturalness_check
        self.topic_consistency = topic_consistency
        self.prompts = prompts or load_prompts(PROMPTS_INI)
    def run(self):
        try:
            model = init_gemini()
            # Build the base prompt
            base_prompt = self.prompts.get("single")
            
            # Add comprehensive feedback handling if provided
            if self.feedback and self.feedback.strip():
                feedback_section = f"""
FEEDBACK ON PREVIOUS QUESTION: {self.feedback.strip()}

IMPORTANT: Generate a completely new question from scratch that addresses the feedback above. 
- If the feedback mentions the question is too easy/hard, adjust the difficulty accordingly
- If the feedback mentions grammar issues, ensure perfect grammar
- If the feedback mentions relevance, make it more relevant to the unit topics
- If the feedback mentions clarity, make it clearer and more specific
- If the feedback mentions engagement, make it more interesting and thought-provoking
- Completely rewrite the question rather than just making minor adjustments
"""
                base_prompt = base_prompt.replace("OUTPUT: only the question text.", f"{feedback_section}\nOUTPUT: only the question text.")
            
            prompt = fmt(
                base_prompt,
                unit_level=self.level,
                level=self.level,  # Add level parameter for prompts.ini template
                modifier=self.modifier,  # Use modifier from advanced features (Upper /Lower /empty)
                unit_title=self.unit_title,
                topics=self.topics_text,
                vocab=", ".join(self.vocab_list) if self.vocab_list else "(none)",
                cefr_tier=self.cefr_tier,
                tier_instructions=get_tier_instructions(self.level, self.cefr_tier),
                quality_validation_instructions=get_quality_validation_instructions(self.quality_validation),
                blooms_taxonomy_instructions=get_blooms_taxonomy_instructions(self.level, self.blooms_level),
                engagement_level_instructions=get_engagement_level_instructions(self.engagement_level),
                academic_background_instructions=get_academic_background_instructions(self.academic_background),
                naturalness_instructions=get_naturalness_instructions(self.naturalness_check),
                strictness_instructions=get_strictness_instructions(self.topic_consistency),
                existing_questions="\n".join([f"- {q}" for q in self.existing])
            )
            
            # Check if formatting worked correctly
            if prompt == base_prompt:
                debug_print(f"WARNING: Single prompt formatting may have failed. Template and result are identical.")
                debug_print(f"Template: {base_prompt[:200]}...")
                debug_print(f"Formatted: {prompt[:200]}...")
            
            resp = model.generate_content(prompt)
            text = getattr(resp, "text", None)
            if not text and hasattr(resp, "candidates") and resp.candidates and len(resp.candidates) > 0:
                candidate = resp.candidates[0]
                if hasattr(candidate, "content") and candidate.content and hasattr(candidate.content, "parts") and candidate.content.parts and len(candidate.content.parts) > 0:
                    text = candidate.content.parts[0].text
            qs = parse_numbered_questions(text or "", max_items=1)
            if not qs:
                line = (text or "").strip().splitlines()[0] if text else ""
                if not line.endswith("?"):
                    line = (line.rstrip(".! ").strip() or "What everyday example fits this topic for you?") + "?"
                qs = [re.sub(r"\s+", " ", line)]
            
            # Log the conversation for debug mode
            log_ai_conversation(
                conversation_type="SINGLE_REGENERATION",
                prompt=prompt,
                response=text,
                error=None
            )
            
            self.ok.emit(self.index, qs[0])
        except Exception as e:
            # Log the error for debug mode
            log_ai_conversation(
                conversation_type="SINGLE_REGENERATION",
                prompt=prompt if 'prompt' in locals() else "Failed to generate prompt",
                response=None,
                error=str(e)
            )
            self.err.emit(f"Gemini single failed: {e}")

# =========================
# DOCX → PDF → PNG (Exact Preview) hardened
# =========================
def _word_com_to_pdf(docx_path, pdf_path):
    if not sys.platform.startswith("win"): raise RuntimeError("Word COM not available on this OS")
    try:
        import pythoncom
        import win32com.client
        from win32com.client import constants
    except Exception as e:
        raise RuntimeError(f"pywin32 not installed: {e}")
    word = None; doc = None
    pythoncom.CoInitialize()
    try:
        word = win32com.client.DispatchEx("Word.Application"); word.Visible = False; word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True, AddToRecentFiles=False, ConfirmConversions=False, Visible=False)
        doc.ExportAsFixedFormat(OutputFileName=os.path.abspath(pdf_path), ExportFormat=constants.wdExportFormatPDF, OpenAfterExport=False,
                                OptimizeFor=constants.wdExportOptimizeForOnScreen, Range=constants.wdExportAllDocument,
                                Item=constants.wdExportDocumentContent, IncludeDocProps=True, KeepIRM=True,
                                CreateBookmarks=constants.wdExportCreateNoBookmarks, DocStructureTags=True,
                                BitmapMissingFonts=True, UseISO19005_1=False)
    finally:
        try:
            if doc is not None: doc.Close(False)
        except Exception: pass
        try:
            if word is not None: word.Quit()
        except Exception: pass
        try:
            pythoncom.CoUninitialize()
        except Exception: pass
    if not os.path.exists(pdf_path): raise RuntimeError("Word COM export did not create PDF")

def _docx2pdf_to_pdf(docx_path, pdf_path):
    try:
        from docx2pdf import convert as docx2pdf_convert
    except Exception as e:
        raise RuntimeError(f"docx2pdf not available: {e}")
    # In windowed apps, stdout/stderr can be None; docx2pdf may write to them
    import io
    old_out, old_err = sys.stdout, sys.stderr
    try:
        if old_out is None: sys.stdout = io.StringIO()
        if old_err is None: sys.stderr = io.StringIO()
        docx2pdf_convert(docx_path, pdf_path)
    finally:
        sys.stdout = old_out
        sys.stderr = old_err
    if not os.path.exists(pdf_path): raise RuntimeError("docx2pdf did not create PDF")

def _soffice_to_pdf(docx_path, pdf_path):
    outdir = os.path.dirname(pdf_path) or "."
    cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path]
    
    # Check if LibreOffice is available
    try:
        subprocess.run(["soffice", "--version"], check=True, capture_output=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        raise RuntimeError(
            "LibreOffice is not installed or not found in PATH.\n\n"
            "To fix this issue:\n"
            "1. Download and install LibreOffice from https://www.libreoffice.org/\n"
            "2. Make sure LibreOffice is added to your system PATH\n"
            "3. Restart the application after installation\n\n"
            "Alternatively, the application will try other PDF conversion methods."
        )
    
    try:
        # Capture stderr to get more detailed error information
        result = subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=60)
    except subprocess.TimeoutExpired:
        raise RuntimeError(
            "LibreOffice conversion timed out (60 seconds).\n\n"
            "This might be due to:\n"
            "1. LibreOffice being busy with another document\n"
            "2. System performance issues\n"
            "3. Large document size\n\n"
            "Please try again or close other LibreOffice instances."
        )
    except subprocess.CalledProcessError as e:
        error_msg = e.stderr.strip() if e.stderr else "Unknown error"
        raise RuntimeError(
            f"LibreOffice conversion failed with error: {error_msg}\n\n"
            "Common solutions:\n"
            "1. Make sure the document is not corrupted\n"
            "2. Check if LibreOffice is up to date\n"
            "3. Try restarting the application\n"
            "4. Ensure you have write permissions to the output directory"
        )
    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice executable not found.\n\n"
            "Please ensure LibreOffice is properly installed and the 'soffice' command is available in your system PATH."
        )
    except Exception as e:
        raise RuntimeError(f"LibreOffice conversion failed: {e}")
    
    produced = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if produced != pdf_path and os.path.exists(produced):
        try: 
            os.replace(produced, pdf_path)
        except Exception:
            import shutil as _sh
            _sh.copyfile(produced, pdf_path)
    if not os.path.exists(pdf_path): 
        raise RuntimeError("LibreOffice did not create PDF file. The conversion process completed but no PDF was generated.")


def _reportlab_to_pdf(docx_path, pdf_path, unit_title="", questions=None):
    """Generate PDF using ReportLab with decorative frames preserved"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, cm
        from reportlab.lib.colors import red, black
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        from reportlab.lib import colors
        from reportlab.graphics.shapes import Drawing, Rect, Line
        from reportlab.graphics import renderPDF
        from reportlab.lib.utils import ImageReader
        from reportlab.pdfgen import canvas
        from reportlab.platypus.flowables import Flowable
    except ImportError:
        raise RuntimeError("ReportLab library not available. Install with: pip install reportlab")
    
    if questions is None:
        questions = []
    
    # Create PDF with decorative elements
    doc_pdf = SimpleDocTemplate(pdf_path, pagesize=A4,
                               rightMargin=2.7*cm, leftMargin=2.7*cm,
                               topMargin=3.6*cm, bottomMargin=2.6*cm)
    
    # Define styles based on template
    styles = getSampleStyleSheet()
    
    # Title style (red, bold, centered) - matching template
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=red,
        alignment=TA_CENTER,
        spaceAfter=30,
        fontName='Helvetica-Bold'
    )
    
    # Discussion style (red, bold)
    discussion_style = ParagraphStyle(
        'CustomDiscussion',
        parent=styles['Normal'],
        fontSize=12,
        textColor=red,
        alignment=TA_LEFT,
        spaceBefore=20,
        spaceAfter=20,
        fontName='Helvetica-Bold'
    )
    
    # Question style with proper indentation
    question_style = ParagraphStyle(
        'CustomQuestion',
        parent=styles['Normal'],
        fontSize=12,
        textColor=black,
        alignment=TA_LEFT,
        spaceBefore=6,
        spaceAfter=12,
        leftIndent=20,
        fontName='Helvetica'
    )
    
    # Build content
    story = []
    
    # Add unit title
    if unit_title:
        story.append(Paragraph(unit_title, title_style))
    else:
        story.append(Paragraph("Unit Title", title_style))
    
    story.append(Spacer(1, 20))
    
    # Add discussion questions
    if questions:
        story.append(Paragraph("Discussion", discussion_style))
        for i, question in enumerate(questions, 1):
            # Format question with proper numbering
            formatted_question = f"{i}. {question}"
            story.append(Paragraph(formatted_question, question_style))
    
    # Custom flowable for decorative frame
    class DecorativeFrame(Flowable):
        def __init__(self, template_image_path=None):
            self.template_image_path = template_image_path
            self.width = A4[0] - 2*cm
            self.height = A4[1] - 2*cm
            
        def draw(self):
            # Draw decorative border
            self.canv.setStrokeColor(red)
            self.canv.setLineWidth(2)
            self.canv.rect(0, 0, self.width, self.height)
            
            # Add corner decorations
            corner_size = 0.5*cm
            self.canv.setFillColor(red)
            # Top-left corner
            self.canv.rect(0, self.height-corner_size, corner_size, corner_size, fill=1)
            # Top-right corner
            self.canv.rect(self.width-corner_size, self.height-corner_size, corner_size, corner_size, fill=1)
            # Bottom-left corner
            self.canv.rect(0, 0, corner_size, corner_size, fill=1)
            # Bottom-right corner
            self.canv.rect(self.width-corner_size, 0, corner_size, corner_size, fill=1)
            
            # Add template background image if available
            if self.template_image_path and os.path.exists(self.template_image_path):
                try:
                    self.canv.drawImage(self.template_image_path, 0, 0, 
                                      width=self.width, height=self.height, 
                                      mask='auto', preserveAspectRatio=True)
                except Exception:
                    pass  # Continue without background image if it fails
    
    # Add decorative frame
    template_image_path = _resource_path("template_base.png")
    story.insert(0, DecorativeFrame(template_image_path))
    
    # Build PDF
    try:
        doc_pdf.build(story)
    except Exception as e:
        raise RuntimeError(f"ReportLab PDF generation failed: {str(e)}")

def check_pdf_conversion_availability():
    """Check which PDF conversion methods are available and return status info"""
    available_methods = []
    missing_methods = []
    
    # Check ReportLab (primary method)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate
        available_methods.append("ReportLab (Self-contained)")
    except Exception:
        missing_methods.append("ReportLab")
    
    # Check docx2pdf (fallback)
    try:
        from docx2pdf import convert
        available_methods.append("docx2pdf library")
    except Exception:
        missing_methods.append("docx2pdf library")
    
    # Check Microsoft Word COM
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Quit()
        available_methods.append("Microsoft Word COM")
    except Exception:
        missing_methods.append("Microsoft Word COM")
    
    # Check LibreOffice
    try:
        subprocess.run(["soffice", "--version"], check=True, capture_output=True)
        available_methods.append("LibreOffice")
    except Exception:
        missing_methods.append("LibreOffice")
    
    
    return available_methods, missing_methods

def convert_docx_to_pdf_with_method(docx_path, pdf_path, method="reportlab", unit_title="", questions=None):
    """Convert DOCX to PDF using a specific method (for debug mode)"""
    if method == "auto":
        # Even in auto mode, use ReportLab as the primary method
        return convert_docx_to_pdf(docx_path, pdf_path, unit_title, questions)
    
    # Try the specific method requested
    try:
        if method == "reportlab":
            _reportlab_to_pdf(docx_path, pdf_path, unit_title, questions)
        elif method == "word_com":
            _word_com_to_pdf(docx_path, pdf_path)
        elif method == "docx2pdf":
            _docx2pdf_to_pdf(docx_path, pdf_path)
        elif method == "libreoffice":
            _soffice_to_pdf(docx_path, pdf_path)
        else:
            raise RuntimeError(f"Unknown PDF conversion method: {method}")
    except Exception as e:
        # If the specific method fails, provide detailed error
        raise RuntimeError(f"Selected PDF method '{method}' failed: {str(e)}")

def convert_docx_to_pdf(docx_path, pdf_path, unit_title="", questions=None):
    """Convert DOCX to PDF using ReportLab (self-contained method)"""
    try: 
        _reportlab_to_pdf(docx_path, pdf_path, unit_title, questions)
        return
    except Exception as e: 
        # If ReportLab fails, try docx2pdf as fallback
        try:
            _docx2pdf_to_pdf(docx_path, pdf_path)
            return
        except Exception as e2:
            # Provide helpful error message for both methods
            error_msg = str(e)
            if "ReportLab library not available" in error_msg:
                raise RuntimeError(
                    "ReportLab library is not installed.\n\n"
                    "To fix this issue:\n"
                    "1. Install ReportLab: pip install reportlab\n"
                    "2. Restart the application after installation\n\n"
                    "This provides self-contained PDF generation without external dependencies."
                )
            elif "docx2pdf not available" in str(e2):
                raise RuntimeError(
                    "Neither ReportLab nor docx2pdf are available.\n\n"
                    "To fix this issue:\n"
                    "1. Install ReportLab: pip install reportlab (recommended)\n"
                    "2. Or install docx2pdf: pip install docx2pdf (requires Word)\n"
                    "3. Restart the application after installation\n\n"
                    "ReportLab provides self-contained PDF generation."
                )
            else:
                raise RuntimeError(
                    f"PDF conversion failed with both methods:\n\n"
                    f"ReportLab error: {error_msg}\n"
                    f"docx2pdf error: {str(e2)}\n\n"
                    "Please install ReportLab for self-contained PDF generation:\n"
                    "pip install reportlab"
                )

class PreviewWorker(QThread):
    done   = pyqtSignal(str, str)   # key, png_path
    failed = pyqtSignal(str)
    def __init__(self, key, unit_title, level, questions, pdf_method="auto"):
        super().__init__(); self.key=key; self.unit_title=unit_title; self.level=level; self.questions=questions[:]
        self._dpr = 1.0  # set by caller for crispness on HiDPI
        self.pdf_method = pdf_method  # PDF conversion method to use
    def _replace_level_and_title(self, doc, idx_disc):
        try: search_range = doc.paragraphs[:max(0, idx_disc)+1]
        except Exception: search_range = doc.paragraphs
        try:
            for p in search_range:
                if "LOOKING BACK" in p.text:
                    for r in p.runs:
                        if "LOOKING BACK" in r.text:
                            r.text = r.text.replace("LOOKING BACK", self.unit_title)
        except Exception: pass
        try:
            for p in search_range:
                for r in p.runs:
                    txt = r.text or ""
                    for tok in LEVEL_TOKENS:
                        if tok in txt: txt = txt.replace(tok, self.level)
                    r.text = txt
        except Exception: pass
    def run(self):
        try:
            from docx import Document
            import fitz
        except Exception as e:
            self.failed.emit(f"Preview deps missing: {e}"); return
        
        if not os.path.exists(TEMPLATE_PATH):
            self.failed.emit("Template not found."); return
        try:
            tmpdir = tempfile.mkdtemp(prefix="efl_prev_")
            tmp_docx = os.path.join(tmpdir, "preview.docx")
            tmp_pdf  = os.path.join(tmpdir, "preview.pdf")
            tmp_png  = os.path.join(tmpdir, "preview.png")

            doc = Document(TEMPLATE_PATH)
            def rm(par):
                p = par._element; p.getparent().remove(p); par._p = par._element = None

            idx_disc=None
            for i,p in enumerate(doc.paragraphs):
                if (p.text or "").strip()=="Discussion": idx_disc=i; break
            if idx_disc is None:
                self.failed.emit('Could not find "Discussion" in template.'); return

            self._replace_level_and_title(doc, idx_disc)

            # remove existing questions, detect numbering style
            number_regex = re.compile(r'^\s*1(\.|\))\s+')
            idx_first_q, prefix_example, style_name = None, "1. ", None
            for i in range(idx_disc+1, len(doc.paragraphs)):
                t = doc.paragraphs[i].text
                if number_regex.match(t or ""):
                    idx_first_q = i; style_name = doc.paragraphs[i].style.name
                    m = re.match(r'^\s*(1(\.|\))\s+)', t)
                    if m: prefix_example = m.group(1)
                    break
            if idx_first_q is None:
                style_name = (doc.paragraphs[idx_disc+1].style.name if idx_disc+1 < len(doc.paragraphs) else "Normal")
                for par in list(doc.paragraphs[idx_disc+1:]): rm(par)
            else:
                for par in list(doc.paragraphs[idx_first_q:]): rm(par)

            dot_or_paren = '.' if '. ' in prefix_example else ')'
            space_part = prefix_example.split(dot_or_paren, 1)[1] if dot_or_paren in prefix_example else ' '
            if not space_part: space_part = ' '

            from docx.shared import Pt
            for i,q in enumerate(self.questions,1):
                para = doc.add_paragraph(style=style_name)
                run = para.add_run(f"{i}{dot_or_paren}{space_part}{q}")
                try: run.font.name = "Tahoma"; run.font.size = Pt(12)
                except Exception: pass

            doc.save(tmp_docx)
            
            # Try to convert DOCX to PDF with better error handling
            try:
                if DEBUG_MODE and self.pdf_method != "reportlab":
                    convert_docx_to_pdf_with_method(tmp_docx, tmp_pdf, self.pdf_method, self.unit_title, self.questions)
                else:
                    convert_docx_to_pdf(tmp_docx, tmp_pdf, self.unit_title, self.questions)
            except Exception as e:
                # Clean up temp files before failing
                try:
                    if os.path.exists(tmp_docx): os.remove(tmp_docx)
                    if os.path.exists(tmp_pdf): os.remove(tmp_pdf)
                    if os.path.exists(tmp_png): os.remove(tmp_png)
                except Exception: pass
                
                # Provide user-friendly error message for preview
                error_msg = str(e)
                if "LibreOffice" in error_msg:
                    self.failed.emit(
                        "Preview generation failed due to PDF conversion issues.\n\n"
                        "The preview requires PDF conversion, but LibreOffice is not available.\n\n"
                        "To fix this:\n"
                        "1. Install LibreOffice from https://www.libreoffice.org/\n"
                        "2. Make sure it's added to your system PATH\n"
                        "3. Restart the application\n\n"
                        "You can still export to DOCX format without preview."
                    )
                elif "All PDF conversion methods failed" in error_msg:
                    self.failed.emit(
                        "Preview generation failed - no PDF conversion method is available.\n\n"
                        "To enable preview:\n"
                        "1. Install LibreOffice (recommended)\n"
                        "2. Or install Microsoft Word (Windows only)\n"
                        "3. Or install docx2pdf library\n\n"
                        "You can still export to DOCX format without preview."
                    )
                else:
                    self.failed.emit(f"Preview generation failed: {error_msg}")

            pdf = fitz.open(tmp_pdf)
            try:
                if pdf.page_count <= 0:
                    raise RuntimeError("PDF has no pages")
                page = pdf[0]
                # Scale by screen DPR for crispness
                mat = fitz.Matrix(2.2 * float(self._dpr), 2.2 * float(self._dpr))
                pix = page.get_pixmap(matrix=mat, alpha=False)
                pix.save(tmp_png)
            finally:
                pdf.close()

            if not os.path.exists(tmp_png) or os.path.getsize(tmp_png) == 0:
                raise RuntimeError("PNG was not produced by renderer")

            self.done.emit(self.key, tmp_png)
        except Exception as e:
            # Clean up temp files on any error
            try:
                if 'tmpdir' in locals():
                    import shutil
                    shutil.rmtree(tmpdir, ignore_errors=True)
            except Exception: pass
            self.failed.emit(str(e))

# =========================
# Preview (empty until exact is requested)
# =========================
class PageWidget(QWidget):
    """
    Canvas that keeps an internal pixmap and exposes `rescale_to_width()`.
    The widget's own size is set to the scaled content size, so the scroll
    area can scroll vertically. SizePolicy is Fixed to avoid layout shrink.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self._pm = QPixmap()
        self._scaled = QPixmap()
        self._scaled_size = QSize(900, 1200)

        sp = QSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        sp.setRetainSizeWhenHidden(True)
        self.setSizePolicy(sp)

    def setPixmap(self, pm: QPixmap):
        # Keep the pixmap’s DPR (do NOT force 1.0)
        self._pm = pm if pm and not pm.isNull() else QPixmap()
        self._scaled = QPixmap()  # invalidate cache
        self.update()

    def rescale_to_fit(self, target_w: int, target_h: int):
        target_w = max(240, int(target_w))
        target_h = max(240, int(target_h))
        if self._pm.isNull():
            # Fallback A4-ish if no pixmap yet; fill by expanding to viewport
            self._scaled = QPixmap()
            self._scaled_size = QSize(target_w, target_h)
            self.resize(self._scaled_size)
            self.update()
            return

        # Work with logical (device-independent) source size
        dpr = max(1.0, float(self._pm.devicePixelRatioF()))
        bw, bh = int(round(self._pm.width() / dpr)), int(round(self._pm.height() / dpr))
        if bw <= 0 or bh <= 0:
            return

        # Fill behavior with minimal margin, no black bars (center-crop to padded rect)
        pad = 8
        avail_w = max(1, target_w - 2*pad)
        avail_h = max(1, target_h - 2*pad)
        src_img = self._pm.toImage()
        scaled_img = src_img.scaled(avail_w, avail_h, Qt.KeepAspectRatioByExpanding, Qt.SmoothTransformation)
        self._scaled = QPixmap.fromImage(scaled_img)

        # Size the widget exactly to the viewport target so it fills the pane
        self._scaled_size = QSize(target_w, target_h)
        self.resize(self._scaled_size)
        self.updateGeometry()
        self.update()

    def sizeHint(self):
        return self._scaled_size

    def minimumSizeHint(self):
        return QSize(200, 280)

    def paintEvent(self, _):
        p = QPainter(self)
        # Fill background to match app palette
        pal = self.palette()
        p.fillRect(self.rect(), pal.window().color())
        if not self._scaled.isNull():
            # Draw inside a slightly padded target area, center-cropped to avoid bars
            pad = 8
            target_rect = self.rect().adjusted(pad, pad, -pad, -pad)
            sw, sh = self._scaled.width(), self._scaled.height()
            tw, th = target_rect.width(), target_rect.height()
            sx = max(0, (sw - tw) // 2)
            sy = max(0, (sh - th) // 2)
            src_rect = QRect(sx, sy, min(tw, sw - sx), min(th, sh - sy))
            p.drawPixmap(target_rect, self._scaled, src_rect)

class BusyOverlay(QWidget):
    """Overlay that does not affect layout and never pushes the page around."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self._text = "Rendering exact…"; self._on = False
        self.setAttribute(Qt.WA_TransparentForMouseEvents, True); self.hide()
    def showMessage(self, text="Rendering exact…"):
        self._text = text; self._on = True
        self.setGeometry(self.parentWidget().rect()); self.show(); self.update()
    def hideMessage(self): self._on = False; self.hide()
    def paintEvent(self, _):
        if not self._on: return
        p = QPainter(self)
        p.fillRect(self.rect(), QColor(0,0,0,100))
        p.setPen(Qt.white); f = p.font(); f.setBold(True); f.setPointSize(12); p.setFont(f)
        p.drawText(self.rect(), Qt.AlignCenter, self._text)

class SimplePreview(QWidget):
    """
    Direct painting preview. Scales to fit width with A4 aspect; allows vertical margins.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self._pm = QPixmap()
        self._scaled = QPixmap()
        self._last_target = QSize(0, 0)
        self._last_scale_width = 0
        self._overlay = BusyOverlay(self)
        sp = QSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        sp.setRetainSizeWhenHidden(True)
        self.setSizePolicy(sp)
        self.setMinimumSize(200, 200)

    def clear(self):
        self._pm = QPixmap(); self.update()

    def queue_exact(self):
        self._overlay.showMessage("Rendering exact…")

    def show_exact(self, png_path: str):
        try:
            if not os.path.exists(png_path):
                raise FileNotFoundError(png_path)
            
            # Clear old content first
            self._pm = QPixmap()
            self._scaled = QPixmap()
            self.update()
            
            pm = QPixmap(png_path)
            # Normalize DPR to avoid Windows HiDPI surprises
            try:
                pm.setDevicePixelRatio(1.0)
            except Exception:
                pass
            if pm.isNull():
                raise RuntimeError("Loaded QPixmap is null")
            self._pm = pm
            # Compute scale once based on current widget size; retry until size is known
            self._rescale_when_ready()
            
            # Also try immediate rescale if widget has size
            if self.width() > 1 and self.height() > 1:
                self._rescale()
            
            # Force multiple updates to ensure display
            self.update()
            self.repaint()
            QCoreApplication.processEvents()
            
            # Ensure widget is visible and properly sized
            self.setVisible(True)
            self.updateGeometry()
            
        except Exception as e:
            QMessageBox.warning(self, "Exact preview error", str(e))
        finally:
            self._overlay.hideMessage()
            # Final update to ensure everything is displayed
            self.update()
            self.repaint()
            QCoreApplication.processEvents()

    def paintEvent(self, _):
        p = QPainter(self)
        pal = self.palette(); p.fillRect(self.rect(), pal.window().color())
        if not self._scaled.isNull():
            pad = 8
            avail_w = max(1, self.width() - 2*pad)
            avail_h = max(1, self.height() - 2*pad)
            x = pad + max(0, (avail_w - self._scaled.width()) // 2)
            # If content taller than available, anchor to top padding; else center vertically
            if self._scaled.height() > avail_h:
                y = pad
            else:
                y = pad + (avail_h - self._scaled.height()) // 2
            p.drawPixmap(x, y, self._scaled)

    def resizeEvent(self, e):
        super().resizeEvent(e)
        # Rescale only when width changes; allow vertical scroll
        if self.width() != self._last_scale_width:
            self._rescale()
        else:
            self.update()

    def _rescale(self):
        if self._pm.isNull():
            self._scaled = QPixmap(); self._last_target = QSize(0, 0); return
        pad = 8
        avail_w = max(1, self.width() - 2*pad)
        # Always fit width to fill the pane at render time; allow vertical overflow or margins
        target_w = avail_w
        target_h = int(round(target_w * 1.4142))
        target = QSize(target_w, target_h)
        
        # Always rescale if we have a new pixmap, even if target size is the same
        # This ensures new content is displayed even when widget size hasn't changed
        img = self._pm.toImage()
        self._scaled = QPixmap.fromImage(img.scaled(target_w, target_h, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        self._last_target = target
        self._last_scale_width = self.width()
        self.update()

    def _rescale_when_ready(self, retries: int = 20):
        if self.width() > 1 and self.height() > 1:
            self._rescale(); return
        if retries <= 0:
            self._rescale(); return
        # Use immediate call for better responsiveness
        QTimer.singleShot(0, lambda: self._rescale_when_ready(retries - 1))

class PreviewPane(QScrollArea):
    """
    Empty preview pane that only displays exact previews when requested.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFrameShape(QScrollArea.NoFrame)
        self.setAlignment(Qt.AlignHCenter | Qt.AlignVCenter)
        self.setWidgetResizable(True)

        self._page = PageWidget()
        self.setWidget(self._page)

        self._overlay = BusyOverlay(self.viewport())
        self._last_target_w = -1

        # Keep scaling synced with viewport size changes
        self.viewport().installEventFilter(self)

    def clear(self):
        self._page.setPixmap(QPixmap())
        self._last_target_w = -1
        self._rescale_to_viewport()

    def queue_exact(self):
        self._overlay.showMessage("Rendering exact…")

    def show_exact(self, png_path: str):
        try:
            if not os.path.exists(png_path):
                raise FileNotFoundError(png_path)
            pm = QPixmap(png_path)
            dpr = max(1.0, float(self.devicePixelRatioF()))
            pm.setDevicePixelRatio(dpr)
            if pm.isNull():
                raise RuntimeError("Loaded QPixmap is null")
            self._page.setPixmap(pm)
            # Ensure page is sized and visible immediately
            try:
                self._page.setVisible(True)
                self._page.resize(self.viewport().size())
            except Exception:
                pass
            # Fit width with aspect; rescale when viewport is ready
            self._rescale_when_ready()
            # Ensure the page widget gets a sane minimum size so it's paintable now
            self._page.setMinimumSize(100, 140)
            # Force a repaint to avoid a stale black viewport
            self._page.update(); self.viewport().update(); self.update()
            # Also schedule a micro deferred repaint to catch late layout
            QTimer.singleShot(0, self._rescale_when_ready)
            QTimer.singleShot(0, self.update)
            QCoreApplication.processEvents()
            # Nudge geometry in case splitter hasn't propagated sizes yet
            self.viewport().setGeometry(self.viewport().geometry())
            self._page.setGeometry(self._page.geometry())
            self._page.update(); self.viewport().update(); self.update()
            # Synchronous repaint
            self.viewport().repaint(); self._page.repaint()
            # Additional delayed stabilization
            QTimer.singleShot(50, self._rescale_when_ready)
            QTimer.singleShot(50, self.update)
        except Exception as e:
            QMessageBox.warning(self, "Exact preview error", str(e))
        finally:
            self._overlay.hideMessage()
            # Ensure overlay removal is reflected immediately
            self._page.update(); self.viewport().update(); self.update()
            QCoreApplication.processEvents()

    # --- sizing helpers ---
    def _effective_viewport_width(self) -> int:
        w = self.viewport().width()
        if w <= 0:
            w = max(self.width(), self.geometry().width(), 0)
        return max(240, w)

    def _rescale_to_viewport(self):
        tw = self._effective_viewport_width()
        # Compute A4 height from width; allow vertical margins, but clamp to viewport height
        th = min(max(240, int(round(tw * 1.4142))), max(240, self.viewport().height()))
        key = (tw, th)
        if key != self._last_target_w:
            self._page.rescale_to_fit(tw, th)
            self._last_target_w = key
            # Ensure layout picks up new page geometry immediately
            self._page.updateGeometry(); self.updateGeometry(); self.viewport().updateGeometry()
            self._page.update(); self.viewport().update(); self.update()
            try:
                # Keep the content in view
                self.ensureWidgetVisible(self._page)
            except Exception:
                pass

    def _rescale_when_ready(self, retries: int = 20):
        vw, vh = self.viewport().width(), self.viewport().height()
        if vw > 1 and vh > 1:
            self._rescale_to_viewport()
            return
        if retries <= 0:
            self._rescale_to_viewport()
            return
        QTimer.singleShot(16, lambda: self._rescale_when_ready(retries - 1))

    # --- Qt overrides ---
    def resizeEvent(self, e):
        super().resizeEvent(e)
        try:
            self._page.resize(self.viewport().size())
        except Exception:
            pass
        self._rescale_to_viewport()

    def showEvent(self, e):
        super().showEvent(e)
        self._rescale_to_viewport()

    def eventFilter(self, obj, ev):
        if obj is self.viewport() and ev.type() == QEvent.Resize:
            self._rescale_to_viewport()
        return super().eventFilter(obj, ev)

# =========================
# Editable rows
# =========================
class EditableLine(QLineEdit):
    edited_confirmed = pyqtSignal(str)
    def __init__(self, text=""):
        super().__init__(text); self.setFont(QFont("Tahoma", 10)); self.setReadOnly(True); self.setFrame(False)
        self._is_real_content = False
        # Style placeholder text to be grayed out, real content to be black
        self.setStyleSheet("""
            QLineEdit[readOnly="true"] {
                color: #666666;
                background-color: #f5f5f5;
            }
            QLineEdit[readOnly="false"] {
                color: #000000;
                background-color: #ffffff;
            }
        """)
    def mouseDoubleClickEvent(self, e): self.setReadOnly(False); self.selectAll(); super().mouseDoubleClickEvent(e)
    def focusOutEvent(self, e):
        super().focusOutEvent(e)
        if not self.isReadOnly(): self.setReadOnly(True); self.edited_confirmed.emit(self.text().strip())
    def keyPressEvent(self, e):
        if e.key() in (Qt.Key_Return, Qt.Key_Enter) and not self.isReadOnly():
            self.setReadOnly(True); self.edited_confirmed.emit(self.text().strip()); return
        super().keyPressEvent(e)
    
    def set_real_content(self, is_real=True):
        """Mark this field as containing real content (not placeholder)"""
        self._is_real_content = is_real
        self._update_style()
    
    def _update_style(self):
        """Update the styling based on content type and read-only state"""
        if self._is_real_content:
            # Real content should be black even when read-only
            self.setStyleSheet("""
                QLineEdit[readOnly="true"] {
                    color: #000000;
                    background-color: #f5f5f5;
                }
                QLineEdit[readOnly="false"] {
                    color: #000000;
                    background-color: #ffffff;
                }
            """)
        else:
            # Placeholder content should be gray when read-only
            self.setStyleSheet("""
                QLineEdit[readOnly="true"] {
                    color: #666666;
                    background-color: #f5f5f5;
                }
                QLineEdit[readOnly="false"] {
                    color: #000000;
                    background-color: #ffffff;
                }
            """)

class QuestionRow(QWidget):
    regen_requested = pyqtSignal(int)
    text_changed    = pyqtSignal(int, str)
    def __init__(self, text, index):
        super().__init__(); self.index=index
        lay = QHBoxLayout(self); lay.setContentsMargins(6,6,6,6); lay.setSpacing(6)
        self.num = QLabel(f"{index}."); self.num.setFont(QFont("Tahoma",10,QFont.Bold)); self.num.setFixedWidth(26)
        self.edit = EditableLine(text); self.edit.edited_confirmed.connect(lambda t: self.text_changed.emit(self.index-1,t))
        # Regenerated indicator (initially hidden)
        self.regen_indicator = QLabel("↻"); self.regen_indicator.setFont(QFont("Tahoma",10,QFont.Bold))
        self.regen_indicator.setStyleSheet("color: #007acc;"); self.regen_indicator.setFixedWidth(16)
        self.regen_indicator.setToolTip("This question was regenerated"); self.regen_indicator.hide()
        btn = QPushButton("↻"); btn.setFixedWidth(34); btn.setToolTip("Regenerate with Gemini")
        btn.setStyleSheet(get_standard_button_style("small"))
        btn.clicked.connect(lambda: self.regen_requested.emit(self.index-1))
        lay.addWidget(self.num,0,Qt.AlignTop); lay.addWidget(self.edit,1); lay.addWidget(self.regen_indicator,0,Qt.AlignTop); lay.addWidget(btn,0,Qt.AlignTop)
    def set_text(self, t): 
        self.edit.setText(t)
        # Mark as real content if it's not a placeholder
        is_real = not (t.startswith("Question ") and t.endswith(" goes here"))
        self.edit.set_real_content(is_real)
    def set_index(self, i): self.index=i; self.num.setText(f"{i}.")
    def mark_as_regenerated(self): self.regen_indicator.show()
    def mark_as_original(self): self.regen_indicator.hide()

# =========================
# Helpers
# =========================
def _safe_open_text(path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return open(path, "a", encoding="utf-8")

def _write_feedback_line(old_question: str, reason: str):
    try:
        from configparser import ConfigParser
        import datetime
        
        # Create or read existing feedback file
        config = ConfigParser()
        if os.path.exists(FEEDBACK_FILE):
            config.read(FEEDBACK_FILE, encoding="utf-8")
        
        # Create feedback section if it doesn't exist
        if not config.has_section("feedback"):
            config.add_section("feedback")
        
        # Generate unique key for this feedback entry
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        key = f"entry_{timestamp}"
        
        # Store feedback data
        reason = (reason or "").strip() or "n/a"
        config.set("feedback", key, f"question={old_question}|reason={reason}")
        
        # Write to file
        with open(FEEDBACK_FILE, "w", encoding="utf-8") as f:
            config.write(f)
    except Exception:
        pass

def _parse_topics_block(text: str, fallback_unit: str):
    unit_title, topics_list, vocab_list = parse_topics_spec(text or "")
    if not unit_title:
        unit_title = fallback_unit
    topics_text = ""
    for t in topics_list:
        topics_text += f"* {t}\n"
    return unit_title, topics_text.strip(), vocab_list

def _ensure_15(qs: list[str]) -> list[str]:
    qs = [re.sub(r"\s+", " ", (q or "").strip()) for q in qs if (q or "").strip()]
    if len(qs) > 15:
        qs = qs[:15]
    while len(qs) < 15:
        qs.append("How could this topic affect an everyday situation at home, work, or school?")
    return qs

# =========================
# About dialog (with updater)
# =========================
class AboutDialog(QDialog):
    def __init__(self, parent=None, auto_check_updates=False):
        super().__init__(parent)
        self.setWindowTitle("EFL Cafe Wizard - About")
        self.setWindowFlags(Qt.Dialog | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
        self.auto_check_updates = auto_check_updates
        self.update_available = False
        
        lay = QVLayout(self)
        title = QLabel(f"<b>{APP_NAME}</b>"); title.setTextFormat(Qt.RichText)
        title.setFont(QFont("Tahoma", 12, QFont.Bold))
        ver   = QLabel(f"Version: <b>{APP_VERSION}</b>"); ver.setTextFormat(Qt.RichText)
        ver.setFont(QFont("Tahoma", 10))
        lay.addWidget(title); lay.addWidget(ver)
        
        # Update status label
        self.update_status = QLabel("Checking for updates...")
        self.update_status.setFont(QFont("Tahoma", 10))
        self.update_status.setStyleSheet("color: #666666;")
        lay.addWidget(self.update_status)
        
        # Debug mode indicator in About dialog
        if DEBUG_MODE:
            debug_label = QLabel("🔧 Debug Mode: DEBUG.txt detected, using prompts.ini from DEBUG folder")
            debug_label.setFont(QFont("Tahoma", 9))
            debug_label.setStyleSheet("color: #ff6b35; background: #fff3e0; padding: 4px 8px; border-radius: 4px; border: 1px solid #ffb74d;")
            lay.addWidget(debug_label)
            
            # Prompts version information
            prompts_version = get_prompts_version()
            if prompts_version:
                prompts_label = QLabel(f"📝 Prompts Version: {prompts_version}")
                prompts_label.setFont(QFont("Tahoma", 9))
                prompts_label.setStyleSheet("color: #2e7d32; background: #e8f5e8; padding: 4px 8px; border-radius: 4px; border: 1px solid #4caf50;")
                lay.addWidget(prompts_label)
            else:
                prompts_label = QLabel("⚠️ Prompts Version: Not found")
                prompts_label.setFont(QFont("Tahoma", 9))
                prompts_label.setStyleSheet("color: #d32f2f; background: #ffebee; padding: 4px 8px; border-radius: 4px; border: 1px solid #f44336;")
                lay.addWidget(prompts_label)
        
        btns = QHBoxLayout()
        self.btn_update = QPushButton("Check for Updates…")
        self.btn_update.setStyleSheet(get_standard_button_style("dialog"))
        btn_ok = QPushButton("Close")
        btn_ok.setStyleSheet(get_standard_button_style("dialog_secondary"))
        self.btn_update.clicked.connect(lambda: check_for_updates(self))
        btn_ok.clicked.connect(self.accept)
        btns.addStretch(1); btns.addWidget(self.btn_update); btns.addWidget(btn_ok)
        lay.addLayout(btns)
        self.resize(350, 120)
        
        # Start background update check if requested
        if auto_check_updates:
            self.start_background_check()
    
    def start_background_check(self):
        """Start background update check and show updater pane if update available"""
        self.update_checker = BackgroundUpdateChecker(self)
        self.update_checker.update_available.connect(self.on_update_available)
        self.update_checker.check_completed.connect(self.on_check_completed)
        self.update_checker.start()
    
    def on_update_available(self, manifest_data):
        """Called when an update is available"""
        self.update_available = True
        latest = manifest_data.get("version", "Unknown")
        notes = manifest_data.get("notes", "")
        
        # Update status label
        self.update_status.setText(f"Update available: v{latest}")
        self.update_status.setStyleSheet("color: #007acc; font-weight: bold;")
        
        # Show update dialog immediately
        self.show_update_dialog(manifest_data)
    
    def on_check_completed(self, update_available):
        """Called when update check is completed"""
        if not update_available:
            self.update_status.setText("You're up to date!")
            self.update_status.setStyleSheet("color: #28a745;")
    
    def show_update_dialog(self, manifest_data):
        """Show the update dialog with the manifest data"""
        latest = manifest_data.get("version", "Unknown")
        notes = manifest_data.get("notes", "")
        
        # Use custom dialog with rich text formatting
        update_dialog = UpdateAvailableDialog(self, latest, APP_VERSION, notes)
        if update_dialog.exec_() == QDialog.Accepted:
            # Proceed with update
            check_for_updates(self, show_updater_pane=True)
    
    def show_feedback(self):
        """Show feedback file location and open it for viewing"""
        try:
            if os.path.exists(FEEDBACK_FILE):
                # Open the feedback file in the default text editor
                if os.name == "nt":  # Windows
                    os.startfile(FEEDBACK_FILE)
                elif os.name == "posix":  # macOS and Linux
                    os.system(f"open '{FEEDBACK_FILE}'" if sys.platform == "darwin" else f"xdg-open '{FEEDBACK_FILE}'")
                QMessageBox.information(self, "Feedback File", 
                    f"Feedback file opened in your default text editor.\n\n"
                    f"Location: {FEEDBACK_FILE}\n\n"
                    f"You can share this file with the developer.")
            else:
                QMessageBox.information(self, "Feedback File", 
                    f"No feedback has been collected yet.\n\n"
                    f"Feedback will be saved to:\n{FEEDBACK_FILE}\n\n"
                    f"Use the regenerate button (↻) on questions to provide feedback.")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Could not open feedback file:\n{e}")

# =========================
# Main App
# =========================
class App(QWidget):
    def __init__(self):
        super().__init__(); self.setWindowTitle(APP_NAME)
        # Set responsive window size based on screen resolution
        self._setup_responsive_window()
        self.level = "B2"; self.unit_title = ""
        self.cefr_tier = "Neutral"  # Upper, Neutral, or Lower (for tier instructions)
        self.modifier = ""  # Empty for neutral, "Upper " or "Lower " for modifiers
        self.prompt_style = "Standard"
        # PDF method selection (debug mode)
        self.selected_pdf_method = "reportlab"  # Default to ReportLab mode
        # Quality control settings
        self.quality_validation = True
        self.blooms_level = "Auto"
        self.engagement_level = "Balanced"
        self.academic_background = False
        self.naturalness_check_value = True
        self.topic_consistency = "Neutral"  # Strict, Neutral, or Free
        self.questions = [
            "Question 1 goes here",
            "Question 2 goes here",
            "Question 3 goes here",
            "Question 4 goes here",
            "Question 5 goes here",
            "Question 6 goes here",
            "Question 7 goes here",
            "Question 8 goes here",
            "Question 9 goes here",
            "Question 10 goes here",
            "Question 11 goes here",
            "Question 12 goes here",
            "Question 13 goes here",
            "Question 14 goes here",
            "Question 15 goes here"
        ]
        self.prompts = load_prompts(PROMPTS_INI); self._workers=set()
        
        # Debug version information
        if DEBUG_MODE:
            debug_print(f"[DEBUG] 🚀 EFL Cafe Wizard v{APP_VERSION} starting...")
            debug_print(f"[DEBUG] 📁 Prompts file: {PROMPTS_INI}")
            debug_print(f"[DEBUG] 🔧 Debug mode: ENABLED")
        
        # Auto-update functionality
        self.background_checker = None
        self.update_timer = None
        self.setup_auto_update()

        self.splitter = QSplitter(Qt.Horizontal)
        # Minimal visual cue for the splitter handle and a thin separator
        self.splitter.setStyleSheet(
            "QSplitter::handle { background: palette(mid); width: 10px; }\n"
            "QSplitter::handle:horizontal { image: none; cursor: col-resize; }"
        )

        # LEFT panel — wider
        left = QWidget(); L = QVLayout(left); L.setContentsMargins(12,12,12,12); L.setSpacing(8)
        hdr = QHBoxLayout(); title = QLabel("EFL Cafe Wizard"); title.setFont(QFont("Tahoma", 12, QFont.Bold))
        
        # Update indicator (initially hidden)
        self.update_indicator = QLabel("🔄 Update Available")
        self.update_indicator.setFont(QFont("Tahoma", 10, QFont.Bold))
        self.update_indicator.setStyleSheet("color: #007acc; background: #e3f2fd; padding: 4px 8px; border-radius: 4px;")
        self.update_indicator.hide()
        self.update_indicator.mousePressEvent = lambda e: self.show_about()
        self.update_indicator.setToolTip("Click to check for updates")
        
        # Debug mode indicator (initially hidden)
        self.debug_indicator = QLabel("🔧 Debug Mode")
        self.debug_indicator.setFont(QFont("Tahoma", 9, QFont.Bold))
        self.debug_indicator.setStyleSheet("color: #ff6b35; background: #fff3e0; padding: 3px 6px; border-radius: 3px; border: 1px solid #ffb74d;")
        self.debug_indicator.hide()
        self.debug_indicator.setToolTip("Debug mode: AI conversation logging active (DEBUG.txt detected)")
        
        self.btn_about = QPushButton("About…"); self.btn_about.clicked.connect(self.show_about)
        self.btn_about.setStyleSheet(get_standard_button_style("outline"))
        self.btn_render = QPushButton("Render Exact Preview"); self.btn_render.clicked.connect(self.render_exact_now)
        self.btn_render.setStyleSheet(get_standard_button_style("primary"))
        hdr.addWidget(title); hdr.addStretch(1); hdr.addWidget(self.debug_indicator); hdr.addWidget(self.update_indicator); hdr.addWidget(self.btn_about); hdr.addWidget(self.btn_render)
        L.addLayout(hdr)

        # Unit title
        title_row = QHBoxLayout(); title_row.addWidget(QLabel("Unit Title:"))
        self.title_edit = QLineEdit(self.unit_title); self.title_edit.setPlaceholderText("Unit Title Goes Here..."); self.title_edit.setFont(QFont("Tahoma", 10))
        self.title_edit.textChanged.connect(self.on_title_changed)
        self.title_edit.setStyleSheet(get_standard_lineedit_style())
        title_row.addWidget(self.title_edit); L.addLayout(title_row)

        # Level selector
        level_row = QHBoxLayout(); level_row.addWidget(QLabel("Level:"))
        self.level_combo = QComboBox(); self.level_combo.setFont(QFont("Tahoma", 9)); self.level_combo.addItems(LEVEL_TOKENS)
        self.level_combo.setCurrentText(self.level); self.level_combo.currentTextChanged.connect(self.on_level_changed)
        self.level_combo.setStyleSheet(get_standard_combo_style())
        level_row.addWidget(self.level_combo)
        
        # Advanced Features Toggle Button (right next to Level)
        self.btn_advanced = QPushButton("⚙ Advanced Features (Experimental)")
        self.btn_advanced.setStyleSheet(get_standard_button_style("minimal"))
        self.btn_advanced.clicked.connect(self.toggle_advanced_features)
        level_row.addWidget(self.btn_advanced)
        level_row.addStretch(1); L.addLayout(level_row)

        # Advanced Features Section (Collapsible)
        self.advanced_frame = QWidget()
        self.advanced_frame.hide()  # Initially hidden
        
        advanced_layout = QVBoxLayout(self.advanced_frame)
        advanced_layout.setContentsMargins(0, 8, 0, 0)
        advanced_layout.setSpacing(6)
        
        # CEFR Tier selector (moved from level row)
        tier_row = QHBoxLayout()
        tier_row.addWidget(QLabel("Tier:"))
        self.cefr_button_group = QButtonGroup()
        self.upper_radio = QRadioButton("Upper")
        self.neutral_tier_radio = QRadioButton("Neutral"); self.neutral_tier_radio.setChecked(True)  # Default to neutral
        self.lower_radio = QRadioButton("Lower")
        self.upper_radio.setStyleSheet(get_standard_radio_style())
        self.neutral_tier_radio.setStyleSheet(get_standard_radio_style())
        self.lower_radio.setStyleSheet(get_standard_radio_style())
        self.cefr_button_group.addButton(self.upper_radio, 0)
        self.cefr_button_group.addButton(self.neutral_tier_radio, 1)
        self.cefr_button_group.addButton(self.lower_radio, 2)
        self.upper_radio.toggled.connect(self.on_cefr_tier_changed)
        self.neutral_tier_radio.toggled.connect(self.on_cefr_tier_changed)
        self.lower_radio.toggled.connect(self.on_cefr_tier_changed)
        tier_row.addWidget(self.upper_radio); tier_row.addWidget(self.neutral_tier_radio); tier_row.addWidget(self.lower_radio)
        tier_row.addStretch(1)
        advanced_layout.addLayout(tier_row)
        
        
        # Prompt Modification Settings
        prompt_mod_row = QHBoxLayout()
        prompt_mod_row.addWidget(QLabel("Prompt Style:"))
        self.prompt_style_combo = QComboBox()
        self.prompt_style_combo.addItems(["Standard", "Concise", "Detailed", "Creative"])
        self.prompt_style_combo.setCurrentText("Standard")
        self.prompt_style_combo.setToolTip("Choose the style of AI prompts:\n• Standard: Balanced approach\n• Concise: Shorter, focused prompts\n• Detailed: Comprehensive instructions\n• Creative: More imaginative approaches")
        self.prompt_style_combo.currentTextChanged.connect(self.on_prompt_style_changed)
        self.prompt_style_combo.setStyleSheet(get_standard_combo_style())
        prompt_mod_row.addWidget(self.prompt_style_combo)
        prompt_mod_row.addStretch(1)
        advanced_layout.addLayout(prompt_mod_row)
        
        # Quality Control Section
        quality_title = QLabel("Quality Control")
        quality_title.setStyleSheet("font-weight: bold; color: #495057; margin-top: 8px;")
        advanced_layout.addWidget(quality_title)
        
        # Question Quality Validation
        quality_validation_row = QHBoxLayout()
        self.quality_validation_check = QCheckBox("Question Quality Validation")
        self.quality_validation_check.setChecked(True)
        self.quality_validation_check.setToolTip("Validates questions for:\n• Clarity and unambiguity\n• Perfect grammar and natural phrasing\n• Cultural sensitivity\n• Appropriate content for the level")
        self.quality_validation_check.toggled.connect(self.on_quality_validation_changed)
        self.quality_validation_check.setStyleSheet(get_standard_checkbox_style())
        quality_validation_row.addWidget(self.quality_validation_check)
        quality_validation_row.addStretch(1)
        advanced_layout.addLayout(quality_validation_row)
        
        # Bloom's Taxonomy Level
        blooms_row = QHBoxLayout()
        blooms_row.addWidget(QLabel("Cognitive Level:"))
        self.blooms_combo = QComboBox()
        self.blooms_combo.addItems(["Auto", "Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create"])
        self.blooms_combo.setCurrentText("Auto")
        self.blooms_combo.setToolTip("Bloom's Taxonomy cognitive levels:\n• Auto: Let AI choose appropriate level\n• Remember: 'What is...?' 'Who...?' 'When...?'\n• Understand: 'Explain why...?' 'Describe...?'\n• Apply: 'How would you use...?' 'Solve...?'\n• Analyze: 'Compare...?' 'What are the differences...?'\n• Evaluate: 'Do you agree...?' 'What's your opinion...?'\n• Create: 'Design...?' 'Invent...?' 'Propose...?'")
        self.blooms_combo.currentTextChanged.connect(self.on_blooms_changed)
        self.blooms_combo.setStyleSheet(get_standard_combo_style())
        blooms_row.addWidget(self.blooms_combo)
        blooms_row.addStretch(1)
        advanced_layout.addLayout(blooms_row)
        
        # Engagement Level
        engagement_row = QHBoxLayout()
        engagement_row.addWidget(QLabel("Engagement Level:"))
        self.engagement_combo = QComboBox()
        self.engagement_combo.addItems(["Balanced", "Low", "Medium", "High", "Very High"])
        self.engagement_combo.setCurrentText("Balanced")
        self.engagement_combo.setToolTip("Question engagement levels:\n• Balanced: Mix of different engagement levels\n• Low: Simple, factual questions\n• Medium: Personal but safe questions\n• High: Thought-provoking, opinion-based\n• Very High: Controversial, challenging topics")
        self.engagement_combo.currentTextChanged.connect(self.on_engagement_changed)
        self.engagement_combo.setStyleSheet(get_standard_combo_style())
        engagement_row.addWidget(self.engagement_combo)
        engagement_row.addStretch(1)
        advanced_layout.addLayout(engagement_row)
        
        # Academic Background Knowledge
        academic_row = QHBoxLayout()
        self.academic_background_check = QCheckBox("Require Academic Background")
        self.academic_background_check.setChecked(False)
        self.academic_background_check.setToolTip("Generate questions that require:\n• Subject-specific knowledge\n• Academic vocabulary\n• Critical thinking skills\n• Research or study background\n\nUseful for advanced students or specialized topics")
        self.academic_background_check.toggled.connect(self.on_academic_background_changed)
        self.academic_background_check.setStyleSheet(get_standard_checkbox_style())
        academic_row.addWidget(self.academic_background_check)
        academic_row.addStretch(1)
        advanced_layout.addLayout(academic_row)
        
        # Naturalness Check
        naturalness_row = QHBoxLayout()
        self.naturalness_check = QCheckBox("Naturalness Check")
        self.naturalness_check.setChecked(True)
        self.naturalness_check.setToolTip("Ensures questions sound natural by:\n• Using native speaker language patterns\n• Appropriate register for the level\n• Natural question flow and structure\n• Avoiding awkward or artificial phrasing")
        self.naturalness_check.toggled.connect(self.on_naturalness_changed)
        self.naturalness_check.setStyleSheet(get_standard_checkbox_style())
        naturalness_row.addWidget(self.naturalness_check)
        naturalness_row.addStretch(1)
        advanced_layout.addLayout(naturalness_row)
        
        
        # Topic Consistency section
        consistency_title = QLabel("Topic Consistency")
        consistency_title.setStyleSheet("font-weight: bold; color: #495057; margin-top: 8px;")
        advanced_layout.addWidget(consistency_title)
        
        consistency_row = QHBoxLayout()
        consistency_row.addWidget(QLabel("Strictness:"))
        self.consistency_button_group = QButtonGroup()
        self.strict_radio = QRadioButton("Strict")
        self.strict_radio.setToolTip("Questions must strictly follow the provided topics\n• No deviation from topic content\n• Questions directly relate to specific topic points\n• Minimal creative interpretation")
        self.neutral_consistency_radio = QRadioButton("Neutral")
        self.neutral_consistency_radio.setChecked(True)  # Default to neutral
        self.neutral_consistency_radio.setToolTip("Balanced approach to topic consistency\n• Questions relate to topics but allow some flexibility\n• Good balance between strict adherence and creativity\n• Recommended for most use cases")
        self.free_radio = QRadioButton("Free")
        self.free_radio.setToolTip("More flexible interpretation of topics\n• Questions can be loosely related to topics\n• Allows creative and tangential connections\n• Good for open-ended discussions")
        self.strict_radio.setStyleSheet(get_standard_radio_style())
        self.neutral_consistency_radio.setStyleSheet(get_standard_radio_style())
        self.free_radio.setStyleSheet(get_standard_radio_style())
        
        self.consistency_button_group.addButton(self.strict_radio, 0)
        self.consistency_button_group.addButton(self.neutral_consistency_radio, 1)
        self.consistency_button_group.addButton(self.free_radio, 2)
        
        self.consistency_button_group.buttonClicked.connect(self.on_consistency_changed)
        
        consistency_row.addWidget(self.strict_radio)
        consistency_row.addWidget(self.neutral_consistency_radio)
        consistency_row.addWidget(self.free_radio)
        consistency_row.addStretch(1)
        advanced_layout.addLayout(consistency_row)
        
        # Wrap advanced features in a scroll area to handle scaling issues
        self.advanced_scroll = QScrollArea()
        self.advanced_scroll.setWidgetResizable(True)
        self.advanced_scroll.setFrameShape(QScrollArea.NoFrame)
        self.advanced_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.advanced_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.advanced_scroll.setMaximumHeight(400)  # Limit height to prevent excessive space usage
        self.advanced_scroll.setWidget(self.advanced_frame)
        self.advanced_scroll.hide()  # Initially hidden by default
        
        L.addWidget(self.advanced_scroll)

        # OCR Upload section
        ocr_section = QHBoxLayout()
        self.btn_upload_ocr = QPushButton("📄 Upload Document for OCR")
        self.btn_upload_ocr.setStyleSheet(get_standard_button_style("success"))
        self.btn_upload_ocr.clicked.connect(self.upload_document_for_ocr)
        self.btn_upload_ocr.setEnabled(OCR_AVAILABLE)
        if not OCR_AVAILABLE:
            self.btn_upload_ocr.setToolTip("OCR not available. Install PIL, pytesseract, and PyMuPDF.")
        
        ocr_section.addWidget(self.btn_upload_ocr)
        ocr_section.addStretch(1)
        L.addLayout(ocr_section)

        # Topics section
        topics_label = QLabel("Topics:")
        topics_label.setStyleSheet("font-weight: bold; color: #495057; margin-top: 8px;")
        L.addWidget(topics_label)
        
        self.topics_box = QTextEdit()
        self.topics_box.setReadOnly(False)
        self.topics_box.setMinimumHeight(120)
        self.topics_box.setStyleSheet(get_standard_textedit_style())
        self.topics_box.setPlaceholderText(
            "Enter topics here. Example:\n"
            "* talk about how we deal with change\n"
            "* talk about past difficulties\n"
            "* talk about daily life in the past\n\n"
            "Or use the 'Upload Document for OCR' button above, or simply drag and drop images/PDFs here to automatically extract topics!"
        )
        L.addWidget(self.topics_box)
        
        # Vocabulary section
        vocab_label = QLabel("Vocabulary (Optional):")
        vocab_label.setStyleSheet("font-weight: bold; color: #495057; margin-top: 8px;")
        L.addWidget(vocab_label)
        
        self.vocab_box = QLineEdit()
        self.vocab_box.setStyleSheet(get_standard_lineedit_style())
        self.vocab_box.setPlaceholderText("Enter vocabulary separated by commas (optional - e.g., accept, adapt, be a step forward, be capable of, resist)")
        L.addWidget(self.vocab_box)

        # Generate button
        gen = QHBoxLayout(); self.btn_generate = QPushButton("Generate 15 with Gemini"); self.btn_generate.clicked.connect(self.generate_with_gemini)
        self.btn_generate.setStyleSheet(get_standard_button_style("primary"))
        gen.addWidget(self.btn_generate); gen.addStretch(1); L.addLayout(gen)

        # Questions list
        self.scroll = QScrollArea(); self.scroll.setWidgetResizable(True)
        self.content = QWidget(); self.content_layout = QVLayout(self.content); self.content_layout.setContentsMargins(6,6,6,6); self.content_layout.setSpacing(6)
        self.row_widgets = []
        for i,q in enumerate(self.questions,1): self._add_row(i,q)
        self.content_layout.addStretch(1); self.scroll.setWidget(self.content)
        L.addWidget(self.scroll, 1)

        # Export buttons
        export_layout = QHBoxLayout()
        self.btn_export_docx = QPushButton("Export to DOCX"); self.btn_export_docx.setFont(QFont("Tahoma", 9)); self.btn_export_docx.clicked.connect(self.export_docx)
        self.btn_export_docx.setStyleSheet(get_standard_button_style("success"))
        self.btn_export_pdf = QPushButton("Export to PDF"); self.btn_export_pdf.setFont(QFont("Tahoma", 9)); self.btn_export_pdf.clicked.connect(self.export_pdf)
        self.btn_export_pdf.setStyleSheet(get_standard_button_style("success"))
        export_layout.addStretch(1)
        export_layout.addWidget(self.btn_export_docx)
        export_layout.addWidget(self.btn_export_pdf)
        L.addLayout(export_layout)
        
        # PDF conversion status indicator (debug mode only)
        if DEBUG_MODE:
            self.pdf_status_label = QLabel()
            self.pdf_status_label.setFont(QFont("Tahoma", 9))
            self.pdf_status_label.setWordWrap(True)
            self.pdf_status_label.setStyleSheet("color: #666; padding: 4px;")
            self._update_pdf_status()
            L.addWidget(self.pdf_status_label)
        else:
            self.pdf_status_label = None
        
        # PDF method selector (debug mode only)
        if DEBUG_MODE:
            self.pdf_method_group = QWidget()
            pdf_method_layout = QVBoxLayout(self.pdf_method_group)
            pdf_method_layout.setContentsMargins(10, 5, 10, 5)
            
            # Title for the selector
            method_title = QLabel("PDF Generation Method (Debug)")
            method_title.setFont(QFont("Tahoma", 10, QFont.Bold))
            method_title.setStyleSheet("color: #007acc; padding: 2px;")
            pdf_method_layout.addWidget(method_title)
            
            # Radio buttons for each method
            self.pdf_method_group_buttons = QButtonGroup()
            self.pdf_methods = [
                ("reportlab", "ReportLab (Self-contained)", True),
                ("docx2pdf", "docx2pdf Library", False),
                ("word_com", "Microsoft Word COM", False),
                ("libreoffice", "LibreOffice", False),
                ("auto", "Auto (Try all methods)", False)
            ]
            
            for method_id, method_name, is_default in self.pdf_methods:
                radio = QRadioButton(method_name)
                radio.setFont(QFont("Tahoma", 9))
                radio.setChecked(is_default)
                radio.setProperty("method_id", method_id)
                radio.setStyleSheet(get_standard_radio_style())
                self.pdf_method_group_buttons.addButton(radio)
                pdf_method_layout.addWidget(radio)
            
            # Connect signal
            self.pdf_method_group_buttons.buttonClicked.connect(self._on_pdf_method_changed)
            
            # Style the group
            self.pdf_method_group.setStyleSheet("""
                QWidget {
                    background-color: #f8f9fa;
                    border: 1px solid #dee2e6;
                    border-radius: 4px;
                    margin: 2px;
                }
                QRadioButton {
                    padding: 2px;
                }
            """)
            
            L.addWidget(self.pdf_method_group)
        else:
            self.pdf_method_group = None

        # RIGHT panel — preview (smaller side)
        right = QWidget(); R = QVLayout(right); R.setContentsMargins(0,0,0,0); R.setSpacing(0)
        # Visual divider between panes
        right.setStyleSheet("background: transparent; border-left: 1px solid palette(midlight);")
        # Wrap preview in a vertical-only scroll area to avoid taskbar overlap
        self.preview = SimplePreview()
        self.preview_scroll = QScrollArea(); self.preview_scroll.setFrameShape(QScrollArea.NoFrame)
        self.preview_scroll.setWidgetResizable(True)
        self.preview_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.preview_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.preview_scroll.setWidget(self.preview)
        R.addWidget(self.preview_scroll)

        self.splitter.addWidget(left); self.splitter.addWidget(right)
        self.splitter.setStretchFactor(0, 3)
        self.splitter.setStretchFactor(1, 2)
        # Start with preview pane hidden
        self.splitter.setSizes([1, 0])
        self.preview.hide()

        main = QHBoxLayout(self); main.setContentsMargins(0,0,0,0); main.addWidget(self.splitter)
        
        # Enable drag and drop for OCR functionality
        self.setAcceptDrops(True)
        
        # Show debug indicator if external prompts.ini is being used
        if DEBUG_MODE:
            self.debug_indicator.show()
            debug_print("DEBUG MODE: External prompts.ini detected and loaded")

    def _setup_responsive_window(self):
        """Set up window size and constraints based on screen resolution"""
        from PyQt5.QtWidgets import QDesktopWidget
        
        # Get available screen dimensions (excluding taskbar)
        screen = QDesktopWidget().availableGeometry()
        screen_width = screen.width()
        screen_height = screen.height()
        
        # Define preferred window size
        preferred_width = 1400
        preferred_height = 900
        
        # Calculate responsive size (max 90% of available screen, min 800x600)
        max_width = int(screen_width * 0.9)
        max_height = int(screen_height * 0.9)
        
        # Use preferred size if it fits, otherwise scale down
        if preferred_width <= max_width and preferred_height <= max_height:
            window_width = preferred_width
            window_height = preferred_height
        else:
            # Scale down proportionally
            scale_factor = min(max_width / preferred_width, max_height / preferred_height)
            window_width = int(preferred_width * scale_factor)
            window_height = int(preferred_height * scale_factor)
        
        # Ensure minimum size
        window_width = max(800, window_width)
        window_height = max(600, window_height)
        
        # Set window size and constraints
        self.resize(window_width, window_height)
        self.setMinimumSize(800, 600)
        self.setMaximumSize(max_width, max_height)
        
        debug_print(f"Available screen: {screen_width}x{screen_height}, Window: {window_width}x{window_height}")

    # ----- helpers -----
    def setup_auto_update(self):
        """Set up automatic update checking"""
        if not AUTO_UPDATE_ENABLED:
            return
        
        # Create background checker
        self.background_checker = BackgroundUpdateChecker(self)
        self.background_checker.update_available.connect(self.on_update_available)
        self.background_checker.check_completed.connect(self.on_update_check_completed)
        
        # Set up periodic checking timer
        self.update_timer = QTimer(self)
        self.update_timer.timeout.connect(self.check_for_updates_background)
        self.update_timer.start(AUTO_UPDATE_INTERVAL_MS)
        
        # Initial check after a delay
        if AUTO_UPDATE_CHECK_ON_STARTUP:
            QTimer.singleShot(5000, self.check_for_updates_background)
    
    def check_for_updates_background(self):
        """Check for updates in the background"""
        if self.background_checker and not self.background_checker.isRunning():
            self.background_checker.start()
    
    def on_update_available(self, manifest_data):
        """Called when an update is available in the background"""
        latest = manifest_data.get("version", "Unknown")
        notes = manifest_data.get("notes", "")
        
        # Show visual indicator
        self.update_indicator.setText(f"🔄 Update v{latest} Available")
        self.update_indicator.show()
        
        # Show custom update dialog with rich text formatting
        update_dialog = UpdateAvailableDialog(self, latest, APP_VERSION, notes)
        if update_dialog.exec_() == QDialog.Accepted:
            check_for_updates(self, show_updater_pane=True)
    
    def on_update_check_completed(self, update_available):
        """Called when background update check is completed"""
        pass  # We don't need to do anything here for background checks
    
    def show_about(self):
        # Show About dialog with auto-update check
        AboutDialog(self, auto_check_updates=True).exec_()

    def _add_row(self, idx, text):
        frame = QFrame(); frame.setFrameShape(QFrame.StyledPanel); fl = QVLayout(frame); fl.setContentsMargins(4,4,4,4)
        row = QuestionRow(text, idx); row.regen_requested.connect(self.regen_single); row.text_changed.connect(self.on_text_edit)
        fl.addWidget(row); self.content_layout.addWidget(frame); self.row_widgets.append(row)

    def _start_exact_render(self):
        # Cancel any existing preview workers to prevent stale results
        for worker in list(self._workers):
            if isinstance(worker, PreviewWorker):
                worker.quit()
                worker.wait(100)  # Wait up to 100ms for cleanup
                self._workers.discard(worker)
        
        key = self._content_key()
        # Pass selected PDF method if in debug mode
        pdf_method = getattr(self, 'selected_pdf_method', 'reportlab') if DEBUG_MODE else 'reportlab'
        worker = PreviewWorker(key, self.unit_title, self.level, self.questions, pdf_method)
        # pass current window DPR to worker (used by PDF→PNG raster)
        try:
            worker._dpr = float(self.devicePixelRatioF())
        except Exception:
            worker._dpr = 1.0
        self._workers.add(worker)

        def on_done(k, png_path):
            try:
                current_key = self._content_key()
                # Check if this is still the most recent request
                if k == current_key:
                    self.preview.show_exact(png_path)
                    # Also update the scroll area to ensure proper display
                    self.preview_scroll.update()
                    self.preview_scroll.repaint()
                    QCoreApplication.processEvents()
                else:
                    # This is a stale result, clean up the temp file
                    try:
                        if os.path.exists(png_path):
                            os.remove(png_path)
                    except Exception:
                        pass
            finally:
                self._workers.discard(worker)

        def on_fail(msg):
            QMessageBox.warning(self, "Exact preview failed", msg)
            self._workers.discard(worker)

        worker.done.connect(on_done); worker.failed.connect(on_fail); worker.start()

    def _content_key(self):
        h = hashlib.sha1(); h.update((self.unit_title or "").encode("utf-8")); h.update((self.level or "").encode("utf-8"))
        for q in self.questions: h.update(b"|"); h.update(q.encode("utf-8"))
        return h.hexdigest()

    # ----- reactive -----
    def on_title_changed(self, txt: str):
        self.unit_title = (txt or "").strip()
        # Clear preview when title changes so it shows updated content
        self.preview.clear()

    def on_level_changed(self, lvl: str):
        self.level = lvl
        # Clear preview when level changes so it shows updated content
        self.preview.clear()

    def on_cefr_tier_changed(self):
        if self.upper_radio.isChecked():
            self.cefr_tier = "Upper"
            self.modifier = "Upper "
        elif self.neutral_tier_radio.isChecked():
            self.cefr_tier = "Neutral"
            self.modifier = ""
        else:  # lower_radio
            self.cefr_tier = "Lower"
            self.modifier = "Lower "
        # Clear preview when tier changes so it shows updated content
        self.preview.clear()


    def on_prompt_style_changed(self):
        """Handle prompt style combo box changes"""
        self.prompt_style = self.prompt_style_combo.currentText()
        # Clear preview when style changes
        self.preview.clear()

    def toggle_advanced_features(self):
        """Toggle the Advanced Features section visibility"""
        if self.advanced_scroll.isVisible():
            self.advanced_scroll.hide()
            self.btn_advanced.setText("⚙ Advanced Features (Experimental)")
        else:
            self.advanced_scroll.show()
            self.btn_advanced.setText("⚙ Advanced Features (Experimental) ▼")

    def on_quality_validation_changed(self):
        """Handle quality validation checkbox changes"""
        self.quality_validation = self.quality_validation_check.isChecked()
        self.preview.clear()

    def on_blooms_changed(self):
        """Handle Bloom's taxonomy level changes"""
        self.blooms_level = self.blooms_combo.currentText()
        self.preview.clear()

    def on_engagement_changed(self):
        """Handle engagement level changes"""
        self.engagement_level = self.engagement_combo.currentText()
        self.preview.clear()

    def on_academic_background_changed(self):
        """Handle academic background requirement changes"""
        self.academic_background = self.academic_background_check.isChecked()
        self.preview.clear()

    def on_naturalness_changed(self):
        """Handle naturalness check changes"""
        self.naturalness_check_value = self.naturalness_check.isChecked()
        self.preview.clear()

    def on_consistency_changed(self):
        """Handle topic consistency changes"""
        if self.strict_radio.isChecked():
            self.topic_consistency = "Strict"
        elif self.neutral_consistency_radio.isChecked():
            self.topic_consistency = "Neutral"
        elif self.free_radio.isChecked():
            self.topic_consistency = "Free"
        self.preview.clear()


    def on_text_edit(self, idx: int, new_text: str):
        if 0 <= idx < len(self.questions):
            self.questions[idx] = new_text.strip()
            # Clear preview when questions change so it shows updated content
            self.preview.clear()

    # ----- Gemini -----
    def generate_with_gemini(self):
        # Check if unit title is empty
        if not self.unit_title or not self.unit_title.strip():
            QMessageBox.warning(self, "Unit Title Required", 
                "Please enter a unit title before generating questions.\n\n"
                "The unit title helps create more relevant and contextual questions.")
            return
        
        # Get topics and vocabulary from separate fields
        topics_text = self.topics_box.toPlainText()
        vocab_text = self.vocab_box.text()
        
        # Parse topics if they contain the old format with vocab
        if topics_text and "vocab:" in topics_text.lower():
            spec = topics_text
            unit_title, topics_text, vocab_list = _parse_topics_block(spec, self.unit_title)
        else:
            # Use separate fields
            unit_title = self.unit_title
            vocab_list = [v.strip() for v in vocab_text.split(',') if v.strip()] if vocab_text else []
        # Update unit title without triggering on_title_changed (which clears preview)
        self.unit_title = unit_title
        self.title_edit.blockSignals(True)
        self.title_edit.setText(self.unit_title)
        self.title_edit.blockSignals(False)

        self.btn_generate.setEnabled(False)
        prog = QProgressDialog("Generating 15 questions…", "Cancel", 0, 0, self)
        prog.setWindowTitle("EFL Cafe Wizard - Generating Questions")
        prog.setWindowFlags(Qt.Dialog | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
        prog.setWindowModality(Qt.ApplicationModal)
        prog.show()

        worker = GeminiBatchWorker(unit_title=self.unit_title, level=self.level,
                                   topics_text=topics_text, vocab_list=vocab_list,
                                   cefr_tier=self.cefr_tier, existing=self.questions, 
                                   modifier=self.modifier, prompt_style=self.prompt_style,
                                   quality_validation=self.quality_validation, blooms_level=self.blooms_level,
                                   engagement_level=self.engagement_level, academic_background=self.academic_background,
                                   naturalness_check=self.naturalness_check_value, topic_consistency=self.topic_consistency, prompts=self.prompts)
        self._workers.add(worker)

        def done(qs):
            try:
                qs = _ensure_15(qs)
                self._apply_batch(qs)
                
            finally:
                prog.close(); self.btn_generate.setEnabled(True); self._workers.discard(worker)

        def fail(msg):
            prog.close(); self.btn_generate.setEnabled(True)
            QMessageBox.critical(self, "Gemini Error", msg); self._workers.discard(worker)

        def cancel():
            worker.terminate()
            worker.wait(3000)  # Wait up to 3 seconds for graceful termination
            prog.close(); self.btn_generate.setEnabled(True); self._workers.discard(worker)

        worker.ok.connect(done); worker.err.connect(fail)
        prog.canceled.connect(cancel)
        worker.start()


    def _apply_batch(self, qs: list[str]):
        self.questions = _ensure_15(qs)
        if len(self.row_widgets) != len(self.questions):
            for w in self.row_widgets: w.setParent(None)
            self.row_widgets.clear()
            for i, q in enumerate(self.questions, 1): self._add_row(i, q)
            self.content_layout.addStretch(1)
        else:
            for i, q in enumerate(self.questions, 1):
                self.row_widgets[i-1].set_text(q); self.row_widgets[i-1].set_index(i)
        # Mark all questions as real content (black text) and original (not regenerated)
        for row in self.row_widgets:
            row.edit.set_real_content(True)
            row.mark_as_original()
        # Auto-trigger exact preview if visible
        try:
            if self.preview.isVisible():
                self.preview.queue_exact()
                # Force a small delay to ensure UI updates are processed
                QTimer.singleShot(100, self._start_exact_render)
        except Exception:
            pass

    def regen_single(self, idx: int):
        if not (0 <= idx < len(self.questions)): return
        
        # Check if unit title is empty
        if not self.unit_title or not self.unit_title.strip():
            QMessageBox.warning(self, "Unit Title Required", 
                "Please enter a unit title before regenerating questions.\n\n"
                "The unit title helps create more relevant and contextual questions.")
            return
        
        old_q = self.questions[idx]
        reason, ok = QInputDialog.getText(self, "Regenerate Question", "What's wrong with this question? (optional):")
        if ok: _write_feedback_line(old_q, reason)

        spec = self.topics_box.toPlainText()
        unit_title, topics_text, vocab_list = _parse_topics_block(spec, self.unit_title)
        # Update unit title without triggering on_title_changed (which clears preview)
        self.unit_title = unit_title
        self.title_edit.blockSignals(True)
        self.title_edit.setText(self.unit_title)
        self.title_edit.blockSignals(False)

        prog = QProgressDialog("Generating a replacement…", "Cancel", 0, 0, self)
        prog.setWindowTitle("EFL Cafe Wizard - Generating Replacement")
        prog.setWindowFlags(Qt.Dialog | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
        prog.setWindowModality(Qt.ApplicationModal)
        prog.show()

        worker = GeminiSingleWorker(index=idx, unit_title=self.unit_title, level=self.level,
                                    topics_text=topics_text, vocab_list=vocab_list,
                                    existing=self.questions, cefr_tier=self.cefr_tier, feedback=reason or "", 
                                    modifier=self.modifier, prompt_style=self.prompt_style,
                                    quality_validation=self.quality_validation, blooms_level=self.blooms_level,
                                    engagement_level=self.engagement_level, academic_background=self.academic_background,
                                    naturalness_check=self.naturalness_check_value, topic_consistency=self.topic_consistency, prompts=self.prompts)
        self._workers.add(worker)

        def done(index, new_q):
            try:
                if new_q and new_q.strip(): self._apply_single(index, new_q.strip())
                else: QMessageBox.warning(self, "No result", "Model returned empty text.")
            finally:
                prog.close(); self._workers.discard(worker)

        def fail(msg):
            prog.close(); QMessageBox.critical(self, "Gemini Error", msg); self._workers.discard(worker)

        def cancel():
            worker.terminate()
            worker.wait(3000)  # Wait up to 3 seconds for graceful termination
            prog.close(); self._workers.discard(worker)

        worker.ok.connect(done); worker.err.connect(fail)
        prog.canceled.connect(cancel)
        worker.start()

    def _apply_single(self, idx: int, new_q: str):
        self.questions[idx] = new_q
        self.row_widgets[idx].set_text(new_q)
        # Mark the regenerated question as real content (black text)
        self.row_widgets[idx].edit.set_real_content(True)
        # Mark as regenerated to show the indicator
        self.row_widgets[idx].mark_as_regenerated()
        # Auto-trigger exact preview if visible
        try:
            if self.preview.isVisible():
                self.preview.queue_exact(); self._start_exact_render()
        except Exception:
            pass

    # ----- Export -----
    def export_docx(self):
        # Check if unit title is empty
        if not self.unit_title or not self.unit_title.strip():
            QMessageBox.warning(self, "Unit Title Required", 
                "Please enter a unit title before exporting.\n\n"
                "The unit title will be used in the exported document.")
            return
        
        path, _ = QFileDialog.getSaveFileName(self, "Export DOCX", f"{self.unit_title}.docx", "Word Document (*.docx)")
        if not path: return
        try:
            from docx import Document
            from docx.shared import Pt
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"python-docx not installed: {e}"); return
        if not os.path.exists(TEMPLATE_PATH):
            QMessageBox.critical(self, "Export Error", f"Template not found:\n{TEMPLATE_PATH}"); return
        try:
            doc = Document(TEMPLATE_PATH)
            def rm(par):
                p = par._element; p.getparent().remove(p); par._p = par._element = None
            try:
                for p in doc.paragraphs:
                    if "LOOKING BACK" in p.text:
                        for r in p.runs:
                            if "LOOKING BACK" in r.text:
                                r.text = r.text.replace("LOOKING BACK", self.unit_title)
                    for r in p.runs:
                        txt = r.text or ""
                        for tok in LEVEL_TOKENS:
                            if tok in txt: r.text = txt.replace(tok, self.level); break
            except Exception: pass
            idx_disc = None
            for i, p in enumerate(doc.paragraphs):
                if (p.text or "").strip() == "Discussion": idx_disc = i; break
            if idx_disc is None:
                QMessageBox.critical(self, "Export Error", 'Could not find "Discussion" in template.'); return
            number_regex = re.compile(r'^\s*1(\.|\))\s+')
            idx_first_q, prefix_example, style_name = None, "1. ", None
            for i in range(idx_disc+1, len(doc.paragraphs)):
                t = doc.paragraphs[i].text
                if number_regex.match(t or ""):
                    idx_first_q = i; style_name = doc.paragraphs[i].style.name
                    m = re.match(r'^\s*(1(\.|\))\s+)', t)
                    if m: prefix_example = m.group(1)
                    break
            if idx_first_q is not None:
                for par in list(doc.paragraphs[idx_first_q:]): rm(par)
            else:
                for par in list(doc.paragraphs[idx_disc+1:]): rm(par)
                style_name = doc.paragraphs[idx_disc+1].style.name if idx_disc+1 < len(doc.paragraphs) else "Normal"
            dot_or_paren = '.' if '. ' in prefix_example else ')'
            space_part = prefix_example.split(dot_or_paren, 1)[1] if dot_or_paren in prefix_example else ' '
            if not space_part: space_part = ' '
            for i, q in enumerate(_ensure_15(self.questions), 1):
                para = doc.add_paragraph(style=style_name)
                run = para.add_run(f"{i}{dot_or_paren}{space_part}{q}")
                try: run.font.name = "Tahoma"; run.font.size = Pt(12)
                except Exception: pass
            doc.save(path)
            QMessageBox.information(self, "Export", f"Saved:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"{e}\n\n{traceback.format_exc()}")

    def _update_pdf_status(self):
        """Update the PDF conversion status indicator (debug mode only)"""
        if not DEBUG_MODE or not hasattr(self, 'pdf_status_label') or self.pdf_status_label is None:
            return
            
        try:
            available, missing = check_pdf_conversion_availability()
            if available:
                if "ReportLab (Self-contained)" in available:
                    status_text = "✓ PDF conversion ready (ReportLab - Self-contained)"
                    self.pdf_status_label.setStyleSheet("color: #28a745; padding: 4px;")
                else:
                    status_text = f"✓ PDF conversion available: {', '.join(available)}"
                    self.pdf_status_label.setStyleSheet("color: #28a745; padding: 4px;")
            else:
                status_text = "⚠ PDF conversion not available. Install ReportLab for self-contained generation."
                self.pdf_status_label.setStyleSheet("color: #dc3545; padding: 4px;")
            
            if missing:
                status_text += f"\nMissing: {', '.join(missing)}"
            
            self.pdf_status_label.setText(status_text)
        except Exception as e:
            self.pdf_status_label.setText(f"Status check failed: {str(e)}")
            self.pdf_status_label.setStyleSheet("color: #dc3545; padding: 4px;")

    def _on_pdf_method_changed(self, button):
        """Handle PDF method selection change (debug mode only)"""
        if not DEBUG_MODE:
            return
            
        method_id = button.property("method_id")
        method_name = button.text()
        
        # Store the selected method for use in PDF conversion
        self.selected_pdf_method = method_id
        
        # Update status to show selected method
        debug_status = f"🔧 Debug: Selected PDF method - {method_name}"
        if hasattr(self, 'pdf_status_label'):
            current_text = self.pdf_status_label.text()
            # Remove any existing debug status line
            lines = current_text.split('\n')
            lines = [line for line in lines if not line.startswith('🔧 Debug:')]
            # Add new debug status
            lines.append(debug_status)
            self.pdf_status_label.setText('\n'.join(lines))

    def export_pdf(self):
        # Check if unit title is empty
        if not self.unit_title or not self.unit_title.strip():
            QMessageBox.warning(self, "Unit Title Required", 
                "Please enter a unit title before exporting.\n\n"
                "The unit title will be used in the exported document.")
            return
        
        path, _ = QFileDialog.getSaveFileName(self, "Export PDF", f"{self.unit_title}.pdf", "PDF Document (*.pdf)")
        if not path: return
        try:
            from docx import Document
            from docx.shared import Pt
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"python-docx not installed: {e}"); return
        if not os.path.exists(TEMPLATE_PATH):
            QMessageBox.critical(self, "Export Error", f"Template not found:\n{TEMPLATE_PATH}"); return
        try:
            # Create temporary DOCX first
            tmp_docx = os.path.join(tempfile.gettempdir(), f"export_{os.getpid()}.docx")
            doc = Document(TEMPLATE_PATH)
            def rm(par):
                p = par._element; p.getparent().remove(p); par._p = par._element = None
            try:
                for p in doc.paragraphs:
                    if "LOOKING BACK" in p.text:
                        for r in p.runs:
                            if "LOOKING BACK" in r.text:
                                r.text = r.text.replace("LOOKING BACK", self.unit_title)
                    for r in p.runs:
                        txt = r.text or ""
                        for tok in LEVEL_TOKENS:
                            if tok in txt: r.text = txt.replace(tok, self.level); break
            except Exception: pass
            idx_disc = None
            for i, p in enumerate(doc.paragraphs):
                if (p.text or "").strip() == "Discussion": idx_disc = i; break
            if idx_disc is None:
                QMessageBox.critical(self, "Export Error", 'Could not find "Discussion" in template.'); return
            number_regex = re.compile(r'^\s*1(\.|\))\s+')
            idx_first_q, prefix_example, style_name = None, "1. ", None
            for i in range(idx_disc+1, len(doc.paragraphs)):
                t = doc.paragraphs[i].text
                if number_regex.match(t or ""):
                    idx_first_q = i; style_name = doc.paragraphs[i].style.name
                    m = re.match(r'^\s*(1(\.|\))\s+)', t)
                    if m: prefix_example = m.group(1)
                    break
            if idx_first_q is not None:
                for par in list(doc.paragraphs[idx_first_q:]): rm(par)
            else:
                for par in list(doc.paragraphs[idx_disc+1:]): rm(par)
                style_name = doc.paragraphs[idx_disc+1].style.name if idx_disc+1 < len(doc.paragraphs) else "Normal"
            dot_or_paren = '.' if '. ' in prefix_example else ')'
            space_part = prefix_example.split(dot_or_paren, 1)[1] if dot_or_paren in prefix_example else ' '
            if not space_part: space_part = ' '
            for i, q in enumerate(_ensure_15(self.questions), 1):
                para = doc.add_paragraph(style=style_name)
                run = para.add_run(f"{i}{dot_or_paren}{space_part}{q}")
                try: run.font.name = "Tahoma"; run.font.size = Pt(12)
                except Exception: pass
            doc.save(tmp_docx)
            
            # Convert DOCX to PDF (use ReportLab with fallback)
            if DEBUG_MODE and hasattr(self, 'selected_pdf_method') and self.selected_pdf_method != "reportlab":
                convert_docx_to_pdf_with_method(tmp_docx, path, self.selected_pdf_method, self.unit_title, _ensure_15(self.questions))
            else:
                convert_docx_to_pdf(tmp_docx, path, self.unit_title, _ensure_15(self.questions))
            
            # Clean up temp file
            try: os.remove(tmp_docx)
            except Exception: pass
            
            QMessageBox.information(self, "Export", f"Saved:\n{path}")
            # Refresh PDF status after successful export (debug mode only)
            if DEBUG_MODE:
                self._update_pdf_status()
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"{e}\n\n{traceback.format_exc()}")

    def render_exact_now(self):
        # Show overlay and kick off exact render immediately
        if not self.preview.isVisible():
            self.preview.show()
            # Give preview reasonable width on reveal
            self.splitter.setSizes([900, max(520, self.width() - 900)])
            # Force layout pass so viewport has correct size before render
            self.splitter.update(); self.layout().activate(); QCoreApplication.processEvents()
            # For SimplePreview just refresh
            self.preview.update(); self.preview.repaint(); QCoreApplication.processEvents()
        else:
            # Preview is already visible, but ensure it's properly refreshed
            self.preview.update(); self.preview.repaint(); QCoreApplication.processEvents()
            # Force layout update to ensure proper display
            self.splitter.update(); self.layout().activate(); QCoreApplication.processEvents()
        
        self.preview.queue_exact()
        self._start_exact_render()

    def upload_document_for_ocr(self):
        """Handle document upload for OCR processing."""
        if not OCR_AVAILABLE:
            QMessageBox.warning(self, "OCR Not Available", 
                "OCR functionality is not available. Please install the required dependencies:\n\n"
                "pip install Pillow pytesseract PyMuPDF")
            return
        
        # File dialog for image and PDF files
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Document for OCR",
            "",
            "Images (*.png *.jpg *.jpeg *.bmp *.tiff *.gif);;PDF Files (*.pdf);;All Files (*)"
        )
        
        if not file_path:
            return
        
        # Determine file type
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.pdf':
            file_type = "pdf"
        elif file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif']:
            file_type = "image"
        else:
            QMessageBox.warning(self, "Unsupported File Type", 
                "Please select an image file (PNG, JPG, JPEG, BMP, TIFF, GIF) or PDF file.")
            return
        
        # Show progress dialog
        self.ocr_progress = QProgressDialog("Processing document...", "Cancel", 0, 0, self)
        self.ocr_progress.setWindowTitle("EFL Cafe Wizard - Processing Document")
        self.ocr_progress.setWindowModality(Qt.WindowModal)
        self.ocr_progress.setWindowFlags(Qt.Dialog | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
        self.ocr_progress.setCancelButton(None)  # Disable cancel for now
        self.ocr_progress.show()
        
        # Start OCR worker
        self.ocr_worker = OCRWorker(file_path, file_type)
        self.ocr_worker.finished.connect(self.on_ocr_finished)
        self.ocr_worker.error.connect(self.on_ocr_error)
        self.ocr_worker.progress.connect(self.on_ocr_progress)
        self.ocr_worker.start()
    
    def on_ocr_progress(self, message):
        """Update OCR progress dialog."""
        if hasattr(self, 'ocr_progress'):
            self.ocr_progress.setLabelText(message)
            QCoreApplication.processEvents()
    
    def on_ocr_finished(self, unit_title, topics_text, vocab_list):
        """Handle successful OCR completion."""
        if hasattr(self, 'ocr_progress'):
            self.ocr_progress.close()
        
        # Update the UI with extracted content
        self.unit_title = unit_title
        self.title_edit.setText(unit_title)
        
        # Update topics box
        self.topics_box.setText(topics_text)
        
        # Update vocabulary box
        if vocab_list:
            vocab_text = ', '.join(vocab_list)
            self.vocab_box.setText(vocab_text)
        else:
            self.vocab_box.setText("")
        
        # Show success message
        QMessageBox.information(self, "OCR Complete", 
            f"Successfully extracted content from document!\n\n"
            f"Unit Title: {unit_title}\n"
            f"Topics: {len(topics_text.split('*')) - 1} topics found\n"
            f"Vocabulary: {len(vocab_list)} words extracted")
        
        # Clear preview since we have new content
        self.preview.clear()
    
    def on_ocr_error(self, error_message):
        """Handle OCR error."""
        if hasattr(self, 'ocr_progress'):
            self.ocr_progress.close()
        
        QMessageBox.critical(self, "OCR Error", f"Failed to process document:\n\n{error_message}")
    
    def dragEnterEvent(self, event):
        """Handle drag enter events for file drops."""
        if event.mimeData().hasUrls():
            # Check if any of the dropped files are supported image or PDF formats
            urls = event.mimeData().urls()
            for url in urls:
                if url.isLocalFile():
                    file_path = url.toLocalFile()
                    file_ext = os.path.splitext(file_path)[1].lower()
                    if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.pdf']:
                        # Add visual feedback by highlighting both fields
                        self.topics_box.setStyleSheet("QTextEdit { font-family: Tahoma; font-size: 11pt; border: 2px dashed #007acc; background-color: #f0f8ff; }")
                        self.vocab_box.setStyleSheet("QLineEdit { font-family: Tahoma; font-size: 11pt; border: 2px dashed #007acc; background-color: #f0f8ff; }")
                        event.acceptProposedAction()
                        return
        event.ignore()
    
    def dragMoveEvent(self, event):
        """Handle drag move events."""
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()
    
    def dragLeaveEvent(self, event):
        """Handle drag leave events to reset visual feedback."""
        # Reset both fields styling
        self.topics_box.setStyleSheet("QTextEdit { font-family: Tahoma; font-size: 11pt; }")
        self.vocab_box.setStyleSheet("QLineEdit { font-family: Tahoma; font-size: 11pt; }")
        super().dragLeaveEvent(event)
    
    def dropEvent(self, event):
        """Handle file drop events for OCR processing."""
        # Reset visual feedback
        self.topics_box.setStyleSheet("QTextEdit { font-family: Tahoma; font-size: 11pt; }")
        self.vocab_box.setStyleSheet("QLineEdit { font-family: Tahoma; font-size: 11pt; }")
        
        if not OCR_AVAILABLE:
            QMessageBox.warning(self, "OCR Not Available", 
                "OCR functionality is not available. Please install the required dependencies:\n\n"
                "pip install Pillow pytesseract PyMuPDF")
            event.ignore()
            return
        
        urls = event.mimeData().urls()
        if not urls:
            event.ignore()
            return
        
        # Process the first valid file
        for url in urls:
            if url.isLocalFile():
                file_path = url.toLocalFile()
                file_ext = os.path.splitext(file_path)[1].lower()
                
                # Check if it's a supported format
                if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif', '.pdf']:
                    # Determine file type
                    if file_ext == '.pdf':
                        file_type = "pdf"
                    else:
                        file_type = "image"
                    
                    # Show progress dialog
                    self.ocr_progress = QProgressDialog("Processing dropped document...", "Cancel", 0, 0, self)
                    self.ocr_progress.setWindowTitle("EFL Cafe Wizard - Processing Document")
                    self.ocr_progress.setWindowModality(Qt.WindowModal)
                    self.ocr_progress.setWindowFlags(Qt.Dialog | Qt.CustomizeWindowHint | Qt.WindowTitleHint)
                    self.ocr_progress.setCancelButton(None)  # Disable cancel for now
                    self.ocr_progress.show()
                    
                    # Start OCR worker
                    self.ocr_worker = OCRWorker(file_path, file_type)
                    self.ocr_worker.finished.connect(self.on_ocr_finished)
                    self.ocr_worker.error.connect(self.on_ocr_error)
                    self.ocr_worker.progress.connect(self.on_ocr_progress)
                    self.ocr_worker.start()
                    
                    event.acceptProposedAction()
                    return
        
        # If no valid files were found
        QMessageBox.information(self, "Unsupported File Type", 
            "Please drop an image file (PNG, JPG, JPEG, BMP, TIFF, GIF) or PDF file.")
        event.ignore()


# =========================
# Entry point
# =========================
def main():
    _enable_win_per_monitor_dpi_awareness()

    # Must be set before constructing QApplication (Qt 5.14+)
    try:
        from PyQt5.QtGui import QGuiApplication
        QGuiApplication.setHighDpiScaleFactorRoundingPolicy(
            getattr(QGuiApplication, "HighDpiScaleFactorRoundingPolicy").PassThrough
        )
    except Exception:
        pass

    try:
        QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
        QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)
    except Exception:
        pass

    app = QApplication(sys.argv)
    # Set Windows AppUserModelID so taskbar uses our icon and groups properly
    try:
        if sys.platform.startswith("win"):
            import ctypes
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("EFL.Cafe.App")
    except Exception:
        pass
    # Set application and window icons
    app_icon = QIcon()
    try:
        icon_path = _resource_path("app.ico")
        if os.path.exists(icon_path):
            app_icon = QIcon(icon_path)
            debug_print(f"Successfully loaded icon from: {icon_path}")
        else:
            debug_print(f"Icon file not found at: {icon_path}")
            # Try alternative paths for development
            alt_paths = ["app.ico", os.path.join(os.getcwd(), "app.ico")]
            for alt_path in alt_paths:
                if os.path.exists(alt_path):
                    app_icon = QIcon(alt_path)
                    debug_print(f"Successfully loaded icon from alternative path: {alt_path}")
                    break
    except Exception as e:
        debug_print(f"Error loading icon: {e}")
    
    # Set the application icon first
    if app_icon and not app_icon.isNull():
        app.setWindowIcon(app_icon)
        debug_print("Application icon set successfully")
    
    # Print debug mode information if active
    debug_mode_info()
    
    # Setup debug folder if in debug mode
    setup_debug_folder()
    
    w = App()
    # Set window icon immediately after creation
    if app_icon and not app_icon.isNull():
        w.setWindowIcon(app_icon)
        # Windows-specific: Set the window icon property
        w.setProperty("windowIcon", app_icon)
        debug_print("Window icon set successfully")
        
        # Windows-specific: Force taskbar icon refresh
        if sys.platform.startswith("win"):
            try:
                # Force window to refresh its icon in taskbar
                w.setWindowFlags(w.windowFlags() | Qt.WindowStaysOnTopHint)
                w.show()
                w.setWindowFlags(w.windowFlags() & ~Qt.WindowStaysOnTopHint)
                w.show()
                debug_print("Windows taskbar icon refresh completed")
            except Exception as e:
                debug_print(f"Windows taskbar icon refresh failed: {e}")
    else:
        debug_print("No valid icon available")
    
    # Center the window on available screen area (excluding taskbar)
    from PyQt5.QtWidgets import QDesktopWidget
    screen = QDesktopWidget().availableGeometry()
    size = w.geometry()
    center_x = screen.x() + (screen.width() - size.width()) // 2
    center_y = screen.y() + (screen.height() - size.height()) // 2
    w.move(center_x, center_y)
    w.show()  # start windowed
    
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()